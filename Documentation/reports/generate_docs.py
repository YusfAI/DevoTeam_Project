"""Génère les deux documents .docx (rapport professionnel + guide technique) au
branding Devoteam, à partir du contenu réel du projet (code, tests, PROGRESS.md).

Script one-shot, à relancer si le contenu doit être régénéré après une évolution
du projet. Nécessite python-docx (voir requirements-dev.txt) — pas une dépendance
runtime de l'application.

Usage : python Documentation/reports/generate_docs.py
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))  # Documentation/reports/
ROOT = os.path.dirname(os.path.dirname(SCRIPT_DIR))
DOC_DIR = SCRIPT_DIR
LOGO = os.path.join(SCRIPT_DIR, "assets", "Dev_logo_rgb.png")
EMAIL_PROOF = os.path.join(SCRIPT_DIR, "assets", "email_proof.png")

CORAL = RGBColor(0xF2, 0x40, 0x5A)
CORAL_HEX = "F2405A"
CORAL_STRONG = RGBColor(0xD4, 0x2A, 0x45)
CORAL_STRONG_HEX = "D42A45"
CORAL_TINT_HEX = "FCE9EC"
CHARCOAL = RGBColor(0x24, 0x24, 0x22)
MUTED = RGBColor(0x52, 0x51, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Calibri"


# ---------- helpers de style ----------

def new_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = CHARCOAL
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.18
    return doc


def _bottom_border(paragraph, hex_color, size="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_cover(doc, doc_label, title, subtitle):
    doc.add_picture(LOGO, width=Inches(2.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_label.upper())
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = CORAL
    run.font.name = FONT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(27)
    run2.font.bold = True
    run2.font.color.rgb = CHARCOAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(13)
    run3.font.color.rgb = MUTED

    for _ in range(6):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mrun = meta.add_run("DevoTeam Dashboard — Projet interne  |  Août 2026  |  Youssef Hmidi")
    mrun.font.size = Pt(10)
    mrun.font.color.rgb = MUTED
    doc.add_page_break()


def add_sommaire(doc, entries):
    add_h1(doc, "Sommaire")
    for entry in entries:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(entry)
        r.font.color.rgb = CHARCOAL
    doc.add_page_break()


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = CORAL
    _bottom_border(p, CORAL_STRONG_HEX)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(13.5)
    run.font.bold = True
    run.font.color.rgb = CHARCOAL
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = CORAL_STRONG
    return p


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_cue(doc, text):
    """Didascalie pour un script oral (ex: [Montrer le chat à l'écran]) — italique,
    ton discret, jamais confondue avec le texte à prononcer."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED
    return p


def add_bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead + " ")
        r.bold = True
        r.font.color.rgb = CORAL_STRONG
    p.add_run(text)
    return p


def add_numbered(doc, text, lead=None):
    p = doc.add_paragraph(style="List Number")
    if lead:
        r = p.add_run(lead + " ")
        r.bold = True
        r.font.color.rgb = CORAL_STRONG
    p.add_run(text)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].paragraphs[0].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        shade_cell(hdr_cells[i], CORAL_HEX)
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if r_i % 2 == 1:
                shade_cell(cells[i], "FAFAF8")
    doc.add_paragraph()
    return table


def add_image_block(doc, path, width_inches, caption=None):
    doc.add_picture(path, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    shade_cell  # noqa: keep import usage grouped
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = CHARCOAL
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F2EE")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "E1E0D9")
        pBdr.append(el)
    pPr.append(pBdr)


# ---------- Document 1 : rapport professionnel ----------

def build_rapport_professionnel():
    doc = new_document()
    add_cover(
        doc,
        "Rapport de projet",
        "DevoTeam Dashboard",
        "Dashboard commercial conversationnel — Rapport professionnel",
    )
    add_sommaire(doc, [
        "Résumé exécutif",
        "Contexte et besoin métier",
        "Solution proposée",
        "Fonctionnalités",
        "Conception et architecture",
        "Stack technique",
        "Sécurité et fiabilité",
        "Preuve de fonctionnement — alertes deadlines",
        "État d'avancement et roadmap",
        "Conclusion",
    ])

    # --- Résumé exécutif ---
    add_h1(doc, "1. Résumé exécutif")
    add_body(doc,
        "DevoTeam Dashboard est une application de tableau de bord commercial "
        "conversationnel : un utilisateur pose une question en français dans un chat "
        "(ex. « budget par pays pour Risk Advisory »), et l'application construit "
        "instantanément un tableau de bord complet qui y répond — plusieurs "
        "graphiques, chiffres clés et tableaux, tous cohérents entre eux, à partir des "
        "données réelles des opportunités commerciales. Aucune compétence SQL ou BI "
        "n'est requise côté utilisateur.")
    add_body(doc,
        "Le projet est aujourd'hui fonctionnel de bout en bout, hébergé localement, "
        "et couvre l'ensemble du cycle : compréhension du langage naturel, génération "
        "de requêtes sécurisées, composition automatique de tableaux de bord "
        "versionnés, et un système d'alerte proactive par email et dans l'interface "
        "pour les opportunités dont l'échéance approche.")
    add_body(doc,
        "Particularité assumée : les données vivent dans un simple Google Sheet, que "
        "les équipes commerciales éditent elles-mêmes. L'application n'a aucune base "
        "de données à installer, sauvegarder ou administrer — un choix qui réduit "
        "fortement le coût d'exploitation pour un portefeuille de cette taille.")

    # --- Contexte et besoin métier ---
    add_h1(doc, "2. Contexte et besoin métier")
    add_body(doc,
        "Les équipes commerciales suivent un grand nombre d'opportunités (appels "
        "d'offres, manifestations d'intérêt, consultations…) réparties sur plusieurs "
        "pays et practices (Digital Transformation, Risk Advisory, Data Management). "
        "Deux problèmes récurrents motivent ce projet :")
    add_bullet(doc,
        "l'analyse de ces données (budget par pays, taux de réussite, répartition "
        "par statut…) nécessite normalement un outil BI ou des requêtes SQL — un frein "
        "pour un usage quotidien rapide par des non-techniciens ;",
        lead="Accès à l'information :")
    add_bullet(doc,
        "une opportunité dont l'échéance approche peut passer inaperçue si personne "
        "ne consulte le dashboard ce jour-là, avec un risque direct de manquer une "
        "date de remise d'offre.",
        lead="Risque de deadline manquée :")
    add_body(doc,
        "Le besoin métier est donc double : donner un accès conversationnel, "
        "immédiat et fiable aux données commerciales, et transformer le suivi des "
        "échéances d'un geste passif (consulter le dashboard) en une alerte active "
        "(le système prévient l'équipe).")

    # --- Solution proposée ---
    add_h1(doc, "3. Solution proposée")
    add_body(doc,
        "L'application répond à ce besoin par un assistant conversationnel branché "
        "directement sur les données commerciales réelles, avec une garantie forte : "
        "le système ne répond jamais par une donnée inventée. Si la demande est "
        "ambiguë (métrique, filtre ou valeur non reconnue), il demande une "
        "clarification plutôt que de deviner un résultat plausible mais faux — un "
        "choix de conception délibéré, central pour la confiance dans les chiffres "
        "affichés à des fins commerciales.")
    add_body(doc,
        "Pour le risque de deadline manquée, une brique d'alerte a été ajoutée : "
        "chaque jour, le système identifie automatiquement les opportunités encore "
        "actives dont l'échéance tombe dans les 7 jours suivants, envoie un email "
        "récapitulatif, et affiche un bandeau d'alerte directement dans le dashboard.")
    add_body(doc,
        "Enfin, plutôt que de livrer un graphique isolé par question, l'application "
        "compose un tableau de bord entier : le chiffre demandé, le même chiffre sous "
        "un autre angle, l'état du pipeline et le détail vérifiable — de quoi non "
        "seulement répondre à la question posée, mais aussi susciter la suivante.")

    # --- Fonctionnalités ---
    add_h1(doc, "4. Fonctionnalités")
    add_h2(doc, "4.1 Chat et compréhension du langage naturel")
    add_bullet(doc, "Questions en français libre, sans syntaxe imposée.")
    add_bullet(doc, "Contexte multi-tour : une question de suivi (« et pour Data Management ? ») "
                     "hérite automatiquement du contexte de la question précédente.")
    add_bullet(doc, "Dates et périodes relatives comprises et calculées de façon fiable "
                     "(« ce mois-ci », « l'année dernière », « le trimestre dernier »…).")
    add_bullet(doc, "Requêtes de comparaison entre plusieurs valeurs (« compare la France et le Maroc »).")
    add_bullet(doc, "Demande de clarification explicite si la question reste ambiguë, plutôt "
                     "qu'une réponse devinée.")
    add_bullet(doc, "L'historique de conversation reste disponible après un rechargement de page — "
                     "une analyse en cours n'est plus perdue par un rafraîchissement accidentel.")

    add_h2(doc, "4.2 Un dashboard complet par question")
    add_body(doc,
        "Une question ne produit pas un graphique isolé mais un tableau de bord "
        "entier : cinq à sept visualisations qui répondent ensemble, et qui partagent "
        "toutes les mêmes filtres — l'utilisateur regarde donc un périmètre cohérent, "
        "jamais des chiffres qui ne se comparent pas entre eux.")
    add_bullet(doc, "Les totaux du périmètre interrogé (budget, nombre d'opportunités, "
                     "montant pondéré) en cartes chiffrées.")
    add_bullet(doc, "Le graphique qui répond directement à la question posée.")
    add_bullet(doc, "Le même chiffre vu sous un autre angle (une dimension complémentaire, "
                     "choisie pour ne jamais dupliquer un axe déjà filtré).")
    add_bullet(doc, "L'état du pipeline commercial sur ce périmètre.")
    add_bullet(doc, "Le détail ligne par ligne, pour vérifier les chiffres.")
    add_body(doc,
        "La page d'accueil, elle, est un tableau de bord fixe répondant aux cinq "
        "questions les plus fréquentes, avec un filtre interactif par practice qui "
        "recalcule instantanément l'ensemble des visualisations.")
    add_bullet(doc, "Neuf types de rendu disponibles : barres, courbes, aires, camemberts, "
                     "cartes chiffrées, tableaux, entonnoir de vente, nuage de points et "
                     "carte de chaleur.")
    add_bullet(doc, "Entonnoir de vente : opportunités par étape du pipeline, dans l'ordre réel "
                     "du cycle de vente — visualise directement où les deals se perdent.")
    add_bullet(doc, "Nuage de points : budget vs probabilité de gain — révèle si les gros "
                     "budgets ont statistiquement plus de chances d'être gagnés.")

    add_h2(doc, "4.3 Tableaux de bord « as code »")
    add_body(doc,
        "Les tableaux de bord ne sont pas configurés dans une interface graphique mais "
        "décrits dans des fichiers texte versionnés (Bruin DAC). Conséquence concrète : "
        "toute évolution d'un tableau de bord passe en revue comme du code, se compare "
        "d'une version à l'autre, et se reproduit à l'identique sur n'importe quel poste.")
    add_bullet(doc, "Les tableaux de bord générés par le chat le sont dans le même format que "
                     "ceux écrits à la main — un tableau de bord produit par l'IA peut donc être "
                     "relu, corrigé et conservé comme n'importe quel autre.")
    add_bullet(doc, "Deux commandes de contrôle : l'une vérifie la structure, l'autre exécute "
                     "réellement chaque requête pour confirmer qu'aucun widget n'est cassé.")

    add_bullet(doc, "Le tableau de bord principal est découpé en cinq sections — offres remises, "
                     "affaires chaudes, santé du portefeuille, pipeline, échéances — accessibles "
                     "par une barre d'onglets. Chacune affiche en sous-titre la question à "
                     "laquelle elle répond, et un bouton ramène au point de départ depuis "
                     "n'importe quel écran.")
    add_bullet(doc, "Quand une section porte déjà la réponse à la question posée, elle s'ouvre "
                     "filtrée comme demandé au lieu qu'un tableau de bord soit créé : "
                     "l'utilisateur reste sur une page qu'il connaît, et le chiffre affiché "
                     "provient d'un widget relu en revue.")

    add_h2(doc, "4.4 Alertes deadlines")
    add_bullet(doc, "Email quotidien automatique récapitulant les opportunités actives à "
                     "échéance ≤ 7 jours, envoyé chaque matin à 8h.")
    add_bullet(doc, "Rappel répété chaque jour tant que l'échéance n'est pas passée — "
                     "aucune opportunité à risque ne peut être oubliée.")
    add_bullet(doc, "Bandeau d'alerte visible directement dans le dashboard, avec le détail "
                     "par opportunité (client, pays, practice, statut, budget, jours restants).")
    add_bullet(doc, "Exclusion automatique des opportunités déjà closes (gagnées, perdues, "
                     "signées, infructueuses…) pour ne pas polluer l'alerte — la même règle "
                     "s'applique désormais à la question « opportunités urgentes » posée "
                     "directement dans le chat, pour une définition cohérente sur tous les canaux.")
    add_bullet(doc, "Donnée « jours restants » recalculée à chaque chargement des données à "
                     "partir de la date réelle du jour, pour ne jamais afficher un chiffre obsolète.")

    add_h2(doc, "4.5 Le Google Sheet comme source de données")
    add_bullet(doc, "Les équipes gèrent les opportunités dans un Google Sheet — un outil qu'elles "
                     "maîtrisent déjà — plutôt que dans une base de données : ajouter ou corriger "
                     "une opportunité ne demande aucune compétence technique.")
    add_bullet(doc, "L'application lit le Sheet directement et le recharge toutes les 15 minutes, "
                     "au démarrage, et à la demande via un bouton dans l'interface. Aucune base de "
                     "données à installer, à sauvegarder ou à administrer.")
    add_bullet(doc, "Chaque ligne est validée indépendamment (dates, montants, statuts) ; une ligne "
                     "invalide est journalisée et ignorée, sans jamais bloquer le chargement des "
                     "autres — c'est ce contrôle qui a révélé douze opportunités au statut erroné, "
                     "passées inaperçues jusque-là.")
    add_bullet(doc, "Les colonnes calculées (mois et année d'échéance, jours restants, montant "
                     "pondéré) sont toujours recalculées à la lecture puis réécrites dans le Sheet : "
                     "elles ne peuvent donc jamais diverger des données saisies.")

    add_h2(doc, "4.6 Branding et expérience utilisateur")
    add_bullet(doc, "Interface habillée aux couleurs et au logo Devoteam.")
    add_bullet(doc, "Micro-interactions et animations (fond animé, transitions, mise en avant "
                     "des nouveaux résultats).")
    add_bullet(doc, "Mode sombre entièrement lisible : deux défauts de contraste corrigés "
                     "(en-tête et texte d'alerte) sans rien changer au mode clair déjà validé.")

    # --- Conception et architecture ---
    add_h1(doc, "5. Conception et architecture")
    add_body(doc,
        "L'architecture suit un principe simple, et non négociable : le LLM ne produit "
        "jamais directement un graphique, une requête ou un chiffre. Il transforme la "
        "question de l'utilisateur en une structure intermédiaire strictement validée "
        "(métrique, dimension, filtres, type de graphique…), et c'est ensuite du code "
        "déterministe qui la traduit en requêtes et en tableau de bord. Cette séparation "
        "garantit que tout ce qui s'affiche provient des données réelles.")
    add_table(doc,
        ["Étape", "Rôle"],
        [
            ["1. Chat utilisateur", "L'utilisateur pose sa question en français dans l'interface."],
            ["2. Compréhension d'intention", "Le LLM (ou un parseur rapide par mots-clés pour les cas "
             "simples) extrait une intention structurée et validée."],
            ["3. Résolution des filtres", "Chaque valeur de filtre (pays, statut…) est vérifiée contre "
             "les données réelles ; en cas de doute, une clarification est demandée."],
            ["4. Composition du tableau de bord", "Des règles déterministes choisissent les cinq à sept "
             "visualisations qui répondent ensemble à la question, toutes filtrées à l'identique."],
            ["5. Génération des requêtes", "Chaque visualisation reçoit sa requête, produite par du code "
             "à partir de l'intention validée — jamais écrite par le LLM."],
            ["6. Affichage", "Le tableau de bord est décrit dans un fichier versionné, puis rendu à "
             "l'écran ; un message texte l'accompagne, calculé directement depuis les données."],
        ])
    add_body(doc,
        "L'alerte deadlines suit le même principe de fiabilité : la vérification des "
        "échéances repose sur un nombre de jours restants recalculé à chaque chargement "
        "des données depuis la date réelle du jour, sur un calendrier automatique "
        "indépendant de toute action utilisateur.")
    add_body(doc,
        "Une seule source de vérité alimente l'ensemble : les données du Google Sheet, "
        "chargées en mémoire. Le moteur de tableaux de bord interroge une copie de "
        "cette même source, régénérée à chaque rafraîchissement — les chiffres du chat "
        "et ceux des tableaux de bord ne peuvent donc pas diverger.", italic=False)

    # --- Stack technique ---
    add_h1(doc, "6. Stack technique")
    add_table(doc,
        ["Brique", "Technologie", "Pourquoi"],
        [
            ["Backend / API", "FastAPI (Python)", "Framework léger, typé, avec validation de données native "
             "(Pydantic) — cohérent avec l'exigence anti-hallucination."],
            ["Source de données", "Google Sheets + pandas", "Les équipes gèrent les opportunités dans un "
             "tableur familier ; l'application le charge en mémoire. Aucune base à installer ni administrer."],
            ["Compréhension du langage", "Google Gemini — gemini-flash-lite-latest", "Sortie JSON "
             "structurée, faible latence, complétée par une validation stricte côté serveur."],
            ["Tableaux de bord", "Bruin DAC", "Tableaux de bord décrits en fichiers versionnés plutôt que "
             "configurés à la souris : relus en revue, reproductibles, 21 types de graphiques."],
            ["Moteur de requête", "DuckDB (fichier local)", "Copie en lecture seule des données, régénérée "
             "à chaque rafraîchissement — le moteur de tableaux de bord n'interroge que du SQL."],
            ["Frontend", "Vite + React", "Interface de chat réactive, build optimisé pour un déploiement "
             "local."],
            ["Planification des tâches", "APScheduler", "Exécution fiable des vérifications quotidiennes "
             "(alertes deadlines, rafraîchissement des données) sans dépendance externe."],
            ["Envoi d'emails", "SMTP Gmail (STARTTLS)", "Solution simple et gratuite pour un usage interne, "
             "sans infrastructure d'envoi supplémentaire à maintenir."],
            ["Tests", "pytest (489 tests)", "Suite automatisée sans dépendance réseau ni données réelles "
             "(mocks), garde-fou contre les régressions."],
        ])
    add_body(doc,
        "Côté données, l'ensemble tient dans une seule feuille de calcul : environ 360 "
        "opportunités commerciales décrites par 18 colonnes (pays, practice, statut, "
        "budget, échéance…). Ce volume tient sans difficulté en mémoire, ce qui rend "
        "toute base de données inutile — un choix qui simplifie radicalement "
        "l'installation et la maintenance, et qui reste valable tant que le portefeuille "
        "ne change pas d'ordre de grandeur.")

    # --- Sécurité et fiabilité ---
    add_h1(doc, "7. Sécurité et fiabilité")
    add_bullet(doc, "Aucune clé API ni identifiant n'est versionné dans le code source "
                     "(fichier .env exclu du dépôt git).")
    add_bullet(doc, "Le LLM ne peut désigner que les colonnes et valeurs d'une liste blanche "
                     "explicite — il n'écrit jamais lui-même de requête.")
    add_bullet(doc, "Toute valeur de filtre non reconnue avec confiance déclenche une demande "
                     "de clarification plutôt qu'une hypothèse silencieuse.")
    add_bullet(doc, "489 tests automatisés couvrent la compréhension du langage, le requêtage des "
                     "données, la génération des tableaux de bord et le système d'alerte.")
    add_bullet(doc, "Le mot de passe d'envoi d'email est un mot de passe d'application dédié "
                     "(jamais le mot de passe principal du compte), également hors du dépôt git.")
    add_bullet(doc, "La clé de compte de service Google (accès au Sheet) est un fichier JSON exclu "
                     "du dépôt (dossier credentials/ entièrement gitignoré) — jamais commité, quel "
                     "que soit son nom.")

    # --- Preuve de fonctionnement ---
    add_h1(doc, "8. Preuve de fonctionnement — alertes deadlines")
    add_body(doc,
        "Capture d'écran ci-dessous : email de rappel automatique réellement reçu, "
        "listant les opportunités actives dont l'échéance tombait dans les 7 jours "
        "suivants au moment de l'envoi (5 et 7 jours restants), avec pays, practice, "
        "statut et budget pour chacune.")
    add_image_block(doc, EMAIL_PROOF, 6.3,
                     caption="Email d'alerte reçu dans la boîte de réception — preuve de bon fonctionnement.")

    # --- État d'avancement et roadmap ---
    add_h1(doc, "10. État d'avancement et roadmap")
    add_body(doc, "L'application est complète et fonctionnelle de bout en bout, en hébergement local. "
                   "Vingt-et-une phases de développement ont été livrées, de la mise en place du "
                   "backend jusqu'à la génération automatique de tableaux de bord complets à partir "
                   "d'une question — en passant par la suppression de la base de données au profit "
                   "d'une lecture directe du Google Sheet.", bold=False)
    add_h2(doc, "Améliorations envisagées (hors périmètre actuel)")
    add_bullet(doc, "Authentification et gestion multi-utilisateurs (restriction par practice/BU).")
    add_bullet(doc, "Épinglage des tableaux de bord générés : permettre à l'utilisateur de conserver "
                     "et renommer un tableau de bord produit par l'IA, pour en faire un véritable "
                     "atelier de création.")
    add_bullet(doc, "Couche sémantique : définir les métriques une seule fois et les réutiliser dans "
                     "tous les tableaux de bord, plutôt que de générer chaque requête isolément.")
    add_bullet(doc, "Suivi analytique des sessions utilisateur.")

    # --- Conclusion ---
    add_h1(doc, "11. Conclusion")
    add_body(doc,
        "DevoTeam Dashboard transforme une problématique classique de BI (accès aux "
        "données commerciales, suivi des échéances) en une expérience conversationnelle "
        "fiable, à la fois simple d'usage pour les équipes commerciales et rigoureuse "
        "dans sa conception technique (validation stricte, tests automatisés, "
        "traçabilité des résultats). Les alertes deadlines par email et dans "
        "l'interface répondent directement au risque métier identifié : ne plus jamais "
        "manquer une échéance commerciale faute d'avoir consulté le dashboard à temps.")
    add_body(doc,
        "Trois partis pris résument le projet. Les données restent là où les équipes "
        "les gèrent déjà — un Google Sheet — ce qui supprime toute base à administrer. "
        "Les tableaux de bord sont du code versionné, donc relisibles et reproductibles "
        "plutôt que configurés à la souris par une seule personne. Et l'intelligence "
        "artificielle sert à comprendre la question, jamais à produire les chiffres : "
        "cette frontière est ce qui rend l'outil utilisable pour de vraies décisions "
        "commerciales.")

    out_path = os.path.join(DOC_DIR, "Rapport_Professionnel_DevoTeam_Dashboard.docx")
    doc.save(out_path)
    return out_path


# ---------- Document 2 : guide technique ----------

def build_guide_technique():
    doc = new_document()
    add_cover(
        doc,
        "Documentation technique",
        "DevoTeam Dashboard",
        "Guide technique détaillé — architecture, code et préparation entretien",
    )
    add_sommaire(doc, [
        "Vue d'ensemble de l'architecture",
        "Stack technique et justifications",
        "Structure du dépôt",
        "Source de données (backend/data_store.py)",
        "Pipeline requête → réponse",
        "Compréhension du langage naturel (backend/llm.py, intent_refiner.py)",
        "Couche de requêtage pandas (backend/db_layer.py, schema_and_whitelist.py)",
        "Tableaux de bord as code (backend/dac_composer.py, sql_builder.py)",
        "Règles métier et assistance (backend/business_rules.py, intent_refiner.py)",
        "Réponses textuelles déterministes (backend/response_builder.py)",
        "Alertes deadlines (backend/alerts.py, scheduler)",
        "Le pont DuckDB (backend/duckdb_export.py)",
        "Frontend (Vite + React)",
        "Sécurité",
        "Tests automatisés",
        "Questions d'entretien probables",
        "Les deux modes d'exécution",
        "Limites connues",
    ])

    # --- Vue d'ensemble ---
    add_h1(doc, "1. Vue d'ensemble de l'architecture")
    add_body(doc,
        "L'application suit un pipeline linéaire, chaque étape gérée par un module "
        "backend dédié. Le principe directeur : le LLM ne produit jamais directement "
        "un graphique, une requête ni un chiffre — il produit uniquement un JSON "
        "intermédiaire strictement validé, que du code déterministe transforme ensuite "
        "en requêtes puis en tableau de bord.")
    add_code_block(doc,
        "Google Sheet (source de vérité)\n"
        "   │  toutes les 15 min (+ démarrage + à la demande)\n"
        "   ▼\n"
        "backend/data_store.py  →  DataFrame pandas (en mémoire)\n"
        "   │                              │\n"
        "   │ duckdb_export.py             └─→ chat & alertes\n"
        "   ▼ (projection lecture seule)\n"
        "dac/data/devoteam.db  ←──────────── interrogé par Bruin DAC\n"
        "\n"
        "Utilisateur (chat React)\n"
        "   │  POST /dashboard { query, previous_intent }\n"
        "   ▼\n"
        "backend/llm.py :: parse_user_query()\n"
        "   │  parseur rapide (mots-clés) OU appel Gemini + validation Pydantic\n"
        "   │  résolution des filtres (fuzzy match) contre les données réelles\n"
        "   ▼  intent = {metric, dimension, filters, chart_type, ...}\n"
        "backend/db_layer.py (pandas)  →  response_builder.py  →  message texte\n"
        "   │\n"
        "backend/dac_composer.py\n"
        "   │  compose 5-7 widgets par règles déterministes\n"
        "   │  sql_builder.py génère le SQL de chacun\n"
        "   ▼  réécrit _principal.yml (travail) + _analyse_<hash>.yml (instantané)\n"
        "Frontend React :: iframe → Bruin DAC (port 8321)\n"
        "   exécute le SQL sur DuckDB et affiche le tableau de bord")
    add_body(doc,
        "En parallèle, un scheduler (APScheduler) exécute deux tâches indépendantes du "
        "chat : le rafraîchissement des données depuis le Sheet toutes les 15 minutes, "
        "et la vérification des échéances proches à 8h (voir section 11).")
    add_body(doc,
        "Trois processus tournent donc en local : l'API FastAPI (port 8000), le serveur "
        "de tableaux de bord DAC (port 8321) et le frontend Vite (port 5173). "
        "scripts/start_dev.bat les lance ensemble.")

    # --- Stack et justifications ---
    add_h1(doc, "2. Stack technique et justifications")
    add_h2(doc, "2.1 De SQLite à MySQL, puis plus de base du tout")
    add_body(doc,
        "Le cadrage initial envisageait SQLite ; le projet est d'abord passé à "
        "MySQL/MariaDB via XAMPP (accès concurrent, vues SQL pré-agrégées par "
        "dimension, pool de connexions). Une fois le Google Sheet devenu l'outil de "
        "saisie des équipes, cette base n'était plus qu'un intermédiaire : les données "
        "y étaient recopiées depuis le Sheet toutes les 15 minutes, pour être relues "
        "juste après. La base a donc été supprimée — le Sheet est désormais lu "
        "directement et chargé dans un DataFrame pandas en mémoire.")
    add_body(doc,
        "Ce que ça a réellement simplifié, au-delà de la suppression de XAMPP des "
        "prérequis : toute la machinerie de vues pré-agrégées a disparu avec elle. La "
        "couche SQL devait choisir la bonne vue selon la dimension demandée, vérifier "
        "que cette vue supportait bien les filtres et la métrique, puis retomber sur "
        "un GROUP BY brut quand ce n'était pas le cas. pandas groupe uniformément, "
        "quelle que soit la combinaison — trois mécanismes de repli en moins.")
    add_body(doc,
        "Le raisonnement tient à l'échelle des données : environ 360 lignes et 18 "
        "colonnes tiennent en mémoire sans effort. Ce choix serait à revoir si le "
        "portefeuille changeait d'ordre de grandeur (dizaines de milliers de lignes), "
        "ou s'il fallait des écritures concurrentes depuis plusieurs postes.")
    add_h2(doc, "2.2 Groq → Google Gemini (migration forcée en cours de projet)")
    add_body(doc,
        "Le mandat initial imposait un modèle open source à faible latence : Groq "
        "servait Llama 3.3 70B avec un `response_format: json_object` garantissant "
        "une sortie JSON syntaxiquement valide, mais pas le respect d'un schéma "
        "strict (contrairement à un mode « structured outputs » complet) — d'où "
        "l'importance de la couche de validation Pydantic + liste blanche placée "
        "juste après (section 6). Ce modèle a depuis été retiré du service par Groq "
        "(erreur 404 modèle introuvable, confirmée avec deux clés API différentes — "
        "pas un problème d'accès mais une dépréciation réelle), cassant toute "
        "question ne passant pas par le chemin rapide par mots-clés. Migration vers "
        "Google Gemini (`gemini-flash-lite-latest` — un alias plutôt qu'une version "
        "épinglée, pour ne pas revivre le même type de panne si Google fait tourner "
        "sa gamme) en gardant volontairement la même architecture (JSON libre + "
        "validation Pydantic + réparation heuristique) plutôt que d'adopter les "
        "sorties structurées strictes de Gemini (vérifiées disponibles à cette "
        "occasion) — changer de fournisseur ET d'architecture en même temps aurait "
        "ajouté un risque inutile sous pression. Deuxième surprise trouvée juste "
        "après, en conditions réelles : le premier modèle essayé (`gemini-flash-"
        "latest`) a un quota gratuit de seulement 5 requêtes/minute, épuisé en "
        "quelques messages de chat — la variante « lite » tient ~16 requêtes/minute "
        "pour une qualité d'extraction équivalente sur les mêmes tests, puisque "
        "c'est le schéma + Pydantic qui garantissent la précision, pas la taille du "
        "modèle. Un quota épuisé (429) et une surcharge transitoire (503, avec deux "
        "tentatives automatiques avant abandon) reçoivent maintenant chacun un "
        "message distinct plutôt que la même erreur générique.")
    add_h2(doc, "2.3 Vega-Lite → Bruin DAC (tableaux de bord as code)")
    add_body(doc,
        "Le rendu était assuré par Vega-Lite : le backend générait une spécification "
        "JSON par question, affichée par un composant React. Ça fonctionnait pour un "
        "graphique isolé, mais chaque type de graphique demandait sa propre logique de "
        "construction — l'entonnoir de vente, par exemple, exigeait de précalculer une "
        "géométrie symétrique en Python parce que Vega-Lite ne sait pas produire cette "
        "forme nativement.")
    add_body(doc,
        "Bruin DAC renverse l'approche : un tableau de bord est un fichier YAML "
        "versionné décrivant des widgets (grille 12 colonnes, 21 types de graphiques, "
        "filtres interactifs), rendu par un binaire dédié. Trois bénéfices concrets : "
        "les tableaux de bord passent en revue comme du code et se comparent d'une "
        "version à l'autre ; l'entonnoir, la carte de chaleur et le nuage de points "
        "sont natifs ; et surtout DAC est explicitement conçu pour que des agents IA "
        "écrivent les tableaux de bord — ce qui correspond exactement à l'usage visé "
        "ici, où le chat en génère un par question.")
    add_body(doc,
        "Contrepartie assumée : DAC sert sa propre interface web (port 8321), qu'on ne "
        "peut pas rendre comme un composant React. L'application l'affiche donc en "
        "iframe, ce qui signifie que la zone tableau de bord n'utilise pas le thème "
        "Devoteam ni le mode sombre du reste de l'interface. C'est le prix payé pour "
        "récupérer un moteur de rendu complet plutôt que de le réécrire.")

    add_h2(doc, "2.4 Pourquoi un fichier DuckDB alors qu'il n'y a plus de base")
    add_body(doc,
        "DAC ne sait interroger que des connexions SQL. Comme les données vivent "
        "désormais dans un DataFrame en mémoire, un pont était nécessaire : "
        "`backend/duckdb_export.py` réécrit ce DataFrame dans un fichier DuckDB local "
        "à chaque rafraîchissement. Ce n'est pas un retour à une base de données — "
        "c'est une projection en lecture seule, régénérée, que rien d'autre ne lit. "
        "pandas reste la source de vérité pour le chat et les alertes, ce qui garantit "
        "que les deux moteurs affichent les mêmes chiffres.")
    add_body(doc,
        "Un piège de concurrence a dû être traité : DuckDB autorise plusieurs lecteurs "
        "OU un seul écrivain, jamais les deux. Une lecture DAC pendant une écriture "
        "échouerait donc. La connexion DAC est déclarée en lecture seule, et l'écriture "
        "réessaie quelques fois à intervalle court — vérifié sous charge avec trois "
        "rafraîchissements concurrents contre vingt-sept requêtes de widgets.")

    add_h2(doc, "2.5 Pourquoi APScheduler plutôt qu'une tâche planifiée externe (cron/Task Scheduler)")
    add_body(doc,
        "Le backend est un processus unique en local (`uvicorn`). APScheduler tourne "
        "dans ce même processus (`BackgroundScheduler`), démarré/arrêté via le cycle "
        "de vie FastAPI (`lifespan`) — aucune dépendance à un ordonnanceur système "
        "externe, ce qui simplifie le déploiement local. Les tâches de démarrage sont "
        "enregistrées comme des jobs « une fois, immédiatement » plutôt qu'exécutées "
        "avant le `yield` : le serveur accepte donc les requêtes dès que la boucle est "
        "prête, sans attendre le premier chargement du Sheet. Limite assumée : si le "
        "processus est arrêté au moment précis d'une exécution planifiée, elle est "
        "manquée — d'où le mécanisme de rattrapage décrit en section 10.")

    # --- Structure du dépôt ---
    add_h1(doc, "3. Structure du dépôt")
    add_table(doc,
        ["Chemin", "Contenu"],
        [
            ["backend/main.py", "Application FastAPI, endpoints, câblage du scheduler."],
            ["backend/data_store.py", "Chargement du Google Sheet vers un DataFrame pandas (source de vérité)."],
            ["backend/llm.py", "Appel LLM, validation Pydantic, résolution des filtres, anti-hallucination."],
            ["backend/intent_refiner.py", "Parseur rapide par mots-clés, dates relatives, affinage de l'intention."],
            ["backend/db_layer.py", "Requêtage pandas du DataFrame (données du message texte)."],
            ["backend/schema_and_whitelist.py", "Source unique de vérité : métriques/dimensions/filtres autorisés."],
            ["backend/sql_builder.py", "Intention validée -> SQL DuckDB. Jamais écrit par le LLM."],
            ["backend/dac_composer.py", "Composition du tableau de bord multi-widgets, écriture du YAML."],
            ["backend/duckdb_export.py", "Projection du DataFrame vers le fichier DuckDB interrogé par DAC."],
            ["backend/response_builder.py", "Messages texte déterministes (jamais générés par le LLM)."],
            ["backend/alerts.py", "Opportunités à échéance proche, envoi de l'email d'alerte."],
            ["backend/business_rules.py", "Règles métier indépendantes de l'affichage (ordre du pipeline)."],
            ["backend/data_quality.py", "Rapport des lignes rejetées et des valeurs manquantes."],
            ["dac/.bruin.yml", "Connexion DuckDB de DAC (aucun identifiant, versionnée volontairement)."],
            ["dac/dashboards/accueil.yml", "Section « Vue d'ensemble commerciale » (11 widgets), produite par scripts/generate_accueil.py, versionnée et relue en revue."],
            ["dac/dashboards/section_*.yml", "Les quatre autres sections du tableau de bord principal (affaires chaudes, santé du portefeuille, pipeline, échéances) — 17 widgets, même générateur."],
            ["backend/overview_match.py", "Décide si une section répond déjà à la question posée, et laquelle ; sa table est prouvée par les tests."],
            ["dac/dashboards/_principal.yml", "Tableau de bord de travail, réécrit par chaque question (éphémère, hors suivi git)."],
            ["dac/dashboards/_analyse_*.yml", "Instantané figé par question, pour les rouvrir (éphémère, hors suivi git)."],
            ["frontend/src/", "Application Vite + React (chat, iframe DAC, hooks, styles)."],
            ["tests/", "Suite pytest — 489 tests, aucune dépendance réseau ni données réelles."],
            ["Documentation/WORKFLOW.md", "Traçage d'une question, du prompt au tableau de bord affiché."],
            ["Documentation/reports/", "Ce guide, le rapport et leur générateur (generate_docs.py)."],
            ["Documentation/planning/", "Brief initial du projet et données sources (hors suivi git du dépôt)."],
        ])

    # --- Source de données ---
    add_h1(doc, "4. Source de données (backend/data_store.py)")
    add_body(doc,
        "Il n'y a plus de base de données. Le Google Sheet est la source de vérité : "
        "`data_store.py` le lit via l'API Google (gspread), valide chaque ligne, et "
        "construit un DataFrame pandas conservé en mémoire. Ce DataFrame est rafraîchi "
        "toutes les 15 minutes, au démarrage, et à la demande via POST /sheets/sync.")

    add_h2(doc, "4.1 Les colonnes")
    add_body(doc,
        "Environ 360 lignes, 18 colonnes. Les quatre dernières sont CALCULÉES à chaque "
        "chargement et jamais lues depuis le Sheet, même si elles y figurent : une "
        "donnée dérivable ne doit jamais pouvoir diverger de ce dont elle dérive.")
    add_table(doc,
        ["Colonne", "Type", "Rôle"],
        [
            ["id", "entier", "Identifiant ; attribué automatiquement (max + 1) et réécrit dans "
             "le Sheet si la ligne n'en a pas."],
            ["country", "texte", "Pays de l'opportunité."],
            ["created_date / deadline", "date", "Dates de création et d'échéance. La deadline est "
             "la source de vérité de toute notion d'urgence."],
            ["practice", "texte", "Digital Transformation / Risk Advisory / Data Management."],
            ["description, buyer, partner", "texte", "Informations descriptives libres."],
            ["opp_type", "texte", "AO, DP, AMI, Consultation, Prospection, Gré à gré, Avant-vente."],
            ["status", "texte", "Statut du cycle de vente — 16 valeurs, dont 7 statuts « clos » "
             "exclus des alertes."],
            ["budget, financial_offer", "décimal", "Montants en euros."],
            ["funding_source", "texte", "Source de financement."],
            ["win_probability", "décimal", "Fraction 0–1 (0.74 = 74 %), jamais stockée déjà "
             "multipliée par 100 — voir response_builder.py."],
            ["deadline_month / deadline_year", "CALCULÉES", "Dérivées de deadline, utilisées comme "
             "dimensions d'analyse."],
            ["days_remaining", "CALCULÉE", "deadline moins la date du jour, recalculée à chaque "
             "chargement — donc jamais obsolète."],
            ["weighted_amount", "CALCULÉE", "financial_offer × win_probability."],
        ])

    add_h2(doc, "4.2 Validation ligne par ligne")
    add_body(doc,
        "Chaque ligne est validée indépendamment : dates acceptées dans deux formats, "
        "nombres à virgule décimale, probabilité tolérée en fraction (0.8) comme en "
        "pourcentage (80), et practice/opp_type/status vérifiés contre la liste blanche "
        "de schema_and_whitelist.py. Une cellule illisible coûte la CELLULE, jamais la "
        "ligne : elle est remplacée par « Non renseigné » et l'opportunité est "
        "conservée avec son budget, son échéance et son client.")
    add_cue(doc,
        "Le rejet de la ligne entière était le choix initial, et il était trop cher : "
        "une opportunité complète disparaissait des totaux à cause d'un statut mal "
        "tapé. « Non renseigné » n'invente rien — il DIT que la donnée manque, reste "
        "visible comme catégorie à part dans les graphiques, et se filtre comme "
        "n'importe quelle autre valeur.")
    add_body(doc,
        "La contrepartie est non négociable : réparer sans tracer reviendrait à "
        "corrompre les données en silence, exactement ce que le rejet servait à "
        "éviter. Chaque remplacement est recensé avec son numéro de ligne, sa colonne "
        "et sa valeur d'origine, exposé par GET /data/quality et par le tableau de "
        "bord « Qualité des données » — qui porte un compteur « Lignes écartées malgré "
        "tout » devant rester à zéro.")
    add_body(doc,
        "Cas particulier de l'échéance : tous les champs dérivés en dépendent "
        "(deadline_month, deadline_year, days_remaining). Sans elle ils restent vides, "
        "et la ligne sort naturellement des analyses temporelles et du filtre "
        "« urgentes » sans cesser de compter dans les totaux — ce qui est précisément "
        "le but de la conserver.")
    add_body(doc,
        "Détail qui a compté en pratique : l'export CSV vers Sheets écrit parfois les "
        "valeurs vides comme le texte littéral « NULL ». Sans normalisation explicite, "
        "un partenaire se serait appelé « NULL ». C'est aussi cette validation qui a "
        "révélé douze opportunités au statut invalide, présentes de longue date et "
        "jamais détectées faute de contrôle à l'écriture. Onze d'entre elles relevaient "
        "d'une décision métier (deux statuts légitimes à ajouter à la liste blanche) ; "
        "les quatre dernières portaient un type d'opportunité dans la colonne statut. "
        "Aucune n'est plus perdue aujourd'hui.")

    add_h2(doc, "4.3 Écriture retour dans le Sheet")
    add_body(doc,
        "Le module écrit dans le Sheet deux choses seulement : l'identifiant des "
        "nouvelles lignes, et les quatre colonnes calculées — pour que l'utilisateur "
        "voie le résultat sans ouvrir l'application. Toutes ces cellules sont envoyées "
        "en un SEUL appel groupé : sur ~360 lignes × 4 colonnes, un appel par cellule "
        "représenterait plus d'un millier d'allers-retours réseau.")
    add_body(doc,
        "Volontairement, aucune suppression : retirer une ligne du Sheet la fait "
        "simplement disparaître du chargement suivant, et rien n'est jamais effacé "
        "ailleurs.")

    add_h2(doc, "4.4 État local (hors Sheet)")
    add_body(doc,
        "Deux éléments seulement vivent hors du Sheet, et aucun n'est une base de "
        "données : `data/scheduler_state.json` retient la date du dernier envoi de "
        "l'alerte quotidienne (anti-doublon, section 10), et le cache des tableaux de "
        "bord déjà générés est un simple dictionnaire en mémoire, vidé à chaque "
        "rafraîchissement des données puisque celles-ci ont pu changer.")

    # --- Pipeline détaillé ---
    add_h1(doc, "5. Pipeline requête → réponse (backend/main.py)")
    add_body(doc, "L'endpoint POST /dashboard orchestre l'ensemble :")
    add_numbered(doc, "appelle parse_user_query(query, previous_intent) — jamais de session serveur, "
                       "le frontend renvoie le dernier intent résolu à chaque nouvelle question.")
    add_numbered(doc, "si l'intention est une conversation générale (salutation, hors sujet) ou "
                       "reste ambiguë, renvoie directement le message de clarification, sans rien interroger.")
    add_numbered(doc, "sinon, exécute build_and_execute_query(intent) sur le DataFrame pandas, puis "
                       "compose le message texte à partir des lignes obtenues.")
    add_numbered(doc, "appelle write_generated_dashboard(query, intent) : le tableau de bord "
                       "multi-widgets est composé puis écrit en YAML, et son nom est renvoyé au "
                       "frontend qui pointe son iframe dessus.")
    add_numbered(doc, "cette génération est enveloppée dans un try/except : si elle échoue, la réponse "
                       "textuelle est renvoyée quand même plutôt que de faire échouer toute la requête.")
    add_numbered(doc, "chaque question est journalisée (logger) avec l'intention extraite, à des fins "
                       "d'audit et d'amélioration du prompt.")

    # --- LLM ---
    add_h1(doc, "6. Compréhension du langage naturel")
    add_h2(doc, "6.1 Deux chemins : parseur rapide vs LLM")
    add_body(doc,
        "try_rule_based_parse() (intent_refiner.py) reconnaît par mots-clés fixes les "
        "cas simples (practice, statut, métrique, dimension). Il n'est utilisé que "
        "s'il n'y a pas de contexte de conversation précédent, et seulement après "
        "passage par _augment_rule_based_result(), qui refuse de lui faire confiance "
        "dès qu'une comparaison ou plusieurs pays sont mentionnés (retour à None → "
        "bascule vers le LLM). Dans tous les autres cas, l'appel Gemini (llm.py) prend "
        "le relais, avec le contexte multi-tour injecté dans le prompt système.")
    add_h2(doc, "6.2 Validation stricte — DashboardIntent (Pydantic)")
    add_body(doc,
        "La sortie JSON du LLM est chargée dans un modèle Pydantic dont le "
        "field_validator sur `filters` rejette toute clé absente de VALID_FILTERS. "
        "Toute erreur de validation renvoie un message de clarification à "
        "l'utilisateur — jamais un objet partiellement rempli avec des valeurs par "
        "défaut devinées.")
    add_h2(doc, "6.3 Résolution des valeurs de filtre — _fuzzy_match()")
    add_body(doc, "Ordre de résolution, dans cet ordre strict :")
    add_numbered(doc, "correspondance exacte dans la liste des valeurs connues ;")
    add_numbered(doc, "alias métier explicite (ex. « gagné » → « Offre gagnée ») ;")
    add_numbered(doc, "sous-chaîne bidirectionnelle ;")
    add_numbered(doc, "correspondance approchée via difflib (seuil 0.72).")
    add_body(doc,
        "Si rien ne correspond avec confiance, IntentUnclear est levée et remonte "
        "comme demande de clarification nommant explicitement la valeur non reconnue "
        "— jamais un filtre silencieusement ignoré ou remplacé.")
    add_h2(doc, "6.4 Dates relatives — calcul déterministe, jamais par le LLM")
    add_body(doc,
        "refine_intent() / _apply_relative_period() (intent_refiner.py) résolvent "
        "des expressions comme « ce mois-ci », « l'année dernière » ou « le trimestre "
        "dernier » par de l'arithmétique de date pure en Python (gestion des passages "
        "d'année incluse). Le prompt système interdit explicitement au LLM de deviner "
        "un calcul de date incertain — un choix qui élimine une classe entière "
        "d'erreurs plausibles-mais-fausses.")

    # --- SQL ---
    add_h1(doc, "7. Couche de requêtage pandas")
    add_body(doc,
        "schema_and_whitelist.py est la source unique de vérité : VALID_METRICS, "
        "VALID_DIMENSIONS, VALID_FILTERS et KNOWN_VALUES bornent ce que le LLM peut "
        "produire. db_layer.py traduit ensuite l'intention validée en opérations "
        "pandas sur le DataFrame en mémoire — ce sont ces données qui alimentent le "
        "message texte affiché dans le chat.")
    add_bullet(doc, "Les filtres deviennent des masques booléens ; un filtre à valeurs multiples "
                     "(comparaison « France et Maroc ») devient un .isin().")
    add_bullet(doc, "Les filtres numériques (range_filters) supportent <, >, <=, >=, = et between.")
    add_bullet(doc, "L'agrégation dépend de la métrique elle-même, pas du champ « aggregation » : "
                     "le budget est toujours sommé, la probabilité de gain toujours moyennée.")
    add_bullet(doc, "La sortie reste une liste de dictionnaires — même contrat qu'à l'époque SQL, "
                     "ce qui a permis de changer de moteur sans toucher response_builder.py.")
    add_body(doc,
        "Simplification obtenue au passage : la version MySQL devait choisir une vue "
        "pré-agrégée selon la dimension, vérifier qu'elle supportait les filtres ET la "
        "métrique demandés, puis retomber sur un calcul groupé quand ce n'était pas le "
        "cas. pandas groupe uniformément quelle que soit la combinaison — ces trois "
        "mécanismes de repli ont disparu avec la base.")

    # --- Tableaux de bord as code ---
    add_h1(doc, "8. Tableaux de bord as code (dac_composer.py, sql_builder.py)")
    add_body(doc,
        "Une question ne produit pas un graphique mais un tableau de bord entier, écrit "
        "en YAML puis rendu par Bruin DAC. Deux modules s'en chargent : dac_composer.py "
        "décide QUELS widgets composent le tableau de bord, sql_builder.py produit la "
        "requête de chacun.")

    add_h2(doc, "8.1 Le LLM ne produit jamais le SQL")
    add_body(doc,
        "C'est la décision structurante de cette partie. Il aurait été tentant de "
        "demander directement le SQL au modèle — c'est ce que font beaucoup d'outils de "
        "« text-to-SQL ». Ça aurait rouvert exactement la classe de bugs que tout le "
        "reste du projet est bâti pour empêcher : colonne inventée, valeur de filtre "
        "hallucinée, agrégation silencieusement fausse.")
    add_body(doc,
        "Le modèle continue donc de ne produire qu'une intention validée contre la "
        "liste blanche, et sql_builder.py la traduit. Le coût de ce choix est nul en "
        "pratique : le Sheet a 18 colonnes fixes, donc la liste blanche décrit déjà "
        "l'intégralité des données interrogeables — il n'existe aucune question "
        "légitime qu'elle empêcherait de poser.")
    add_bullet(doc, "Les noms de colonnes proviennent toujours de la liste blanche, jamais de "
                     "texte libre : une métrique ou une dimension inconnue est ignorée et remplacée "
                     "par le défaut, jamais insérée dans la requête.")
    add_bullet(doc, "Les valeurs de filtre sont échappées (les apostrophes doublées) — les vraies "
                     "données en contiennent (« Côte d'Ivoire », « Complément d'information »), et "
                     "sans ce traitement la requête serait syntaxiquement cassée.")

    add_h2(doc, "8.2 Composition par règles déterministes")
    add_body(doc,
        "L'intention donne l'angle principal ; des règles fixes l'entourent de widgets "
        "complémentaires qui partagent TOUS ses filtres — sans quoi le tableau de bord "
        "mélangerait des chiffres qui ne se comparent pas.")
    add_table(doc,
        ["Widget", "Rôle"],
        [
            ["Totaux du périmètre", "La métrique demandée, le nombre d'opportunités et le montant "
             "pondéré, en cartes chiffrées."],
            ["Graphique principal", "Répond directement à la question posée."],
            ["Angle complémentaire", "La même métrique sur une autre dimension."],
            ["Pipeline commercial", "L'entonnoir des statuts sur ce périmètre."],
            ["Détail", "Les opportunités ligne par ligne, pour vérifier les chiffres."],
        ])
    add_body(doc,
        "Pourquoi des règles plutôt qu'un second appel au LLM : la pertinence de ces "
        "widgets se déduit entièrement de l'intention principale, et un appel "
        "supplémentaire par question consommerait le quota gratuit du modèle (~16 "
        "requêtes/minute) pour un gain nul.")

    add_h2(doc, "8.3 Deux défauts trouvés en exécutant, pas en relisant")
    add_bullet(doc, "Le widget complémentaire pouvait reprendre une dimension DÉJÀ FIGÉE par un "
                     "filtre : demander « budget par pays pour Risk Advisory » puis afficher « budget "
                     "par practice » donnait un graphique à une seule barre, sans aucune information. "
                     "La dimension complémentaire évite désormais toute dimension épinglée par un "
                     "filtre — défaut invisible en lecture du code, évident dès la première exécution.")
    add_bullet(doc, "PyYAML repliait le SQL en style quoté, produisant un YAML correct mais "
                     "illisible — ce qui ruinait l'intérêt même du « dashboard as code », dont toute "
                     "la valeur est d'être relu en revue. Un dumper en bloc littéral le garde lisible "
                     "tel qu'écrit.")

    add_h2(doc, "8.4 L'entonnoir de vente")
    add_body(doc,
        "Cas le plus instructif. Les étapes sont ordonnées par l'ordre RÉEL du cycle de "
        "vente, via un CASE explicite dans la requête — jamais par volume, qui "
        "donnerait un entonnoir ne racontant rien du parcours d'une opportunité. Les "
        "statuts de sortie (perdu, NO GO, infructueux…) en sont exclus : ce sont des "
        "sorties du pipeline, pas des étapes que chaque opportunité traverse.")
    add_body(doc,
        "À noter : avec Vega-Lite, cette forme exigeait de précalculer en Python une "
        "géométrie symétrique (bornes gauche/droite centrées sur zéro) car la "
        "bibliothèque ne la produit pas nativement. DAC ayant un type funnel natif, ce "
        "contournement a disparu avec la migration.")

    # --- Réponses texte ---

    # --- Règles métier et assistance ---
    add_h1(doc, "9. Règles métier et assistance (business_rules.py, intent_refiner.py)")
    add_body(doc,
        "Trois familles de règles décident de ce qui est compté, de la forme retenue "
        "pour le montrer, et de ce que devient une demande d'ajustement. Elles vivent "
        "hors du modèle de langage : ce sont des décisions métier, elles doivent être "
        "relisibles, testables et identiques d'une exécution à l'autre.")

    add_h2(doc, "9.1 Les affaires perdues sortent des chiffres par défaut")
    add_body(doc,
        "Un « budget total » de 170,0 M DT additionnait 66,1 M DT d'affaires définitivement "
        "mortes — Offre perdue, Infructueux, NO GO, Hors scope, Non shortlisté. Le "
        "chiffre était exact et inutilisable : personne ne décide sur un portefeuille "
        "qui compte ses échecs comme du potentiel. La liste LOST_STATUSES est donc "
        "retirée par défaut de toutes les métriques et de tous les graphiques.")
    add_cue(doc,
        "Le point délicat est la SYMÉTRIE. Deux moteurs répondent à la même question : "
        "le SQL des widgets (sql_builder) et le pandas du message texte (db_layer). Si "
        "un seul appliquait la règle, le graphique et la phrase juste au-dessus "
        "annonceraient deux totaux différents. Les deux l'appliquent, et l'égalité a "
        "été vérifiée sur les vraies données — 98 320 001 € des deux côtés avant "
        "récupération des lignes rejetées.")
    add_body(doc,
        "C'est un défaut, pas une censure. Dès que la question filtre elle-même sur un "
        "statut, l'exclusion est levée : « liste des offres perdues » doit répondre, "
        "pas renvoyer un tableau vide. Une simple répartition PAR statut ne la lève "
        "pas, sans quoi les barres ne se réconcilieraient plus avec le total affiché "
        "au-dessus d'elles.")

    add_h2(doc, "9.2 L'entonnoir, reconstruit depuis un instantané")
    add_body(doc,
        "FUNNEL_STAGE_ORDER décrit les 12 étapes du pipeline, de Lead à Offre signée. "
        "La subtilité est que le statut est l'état COURANT d'une opportunité, pas son "
        "historique : compter les opportunités par statut ne décroît pas et ne forme "
        "donc aucun entonnoir. La reconstruction cumule depuis la fin — « ayant atteint "
        "au moins cette étape » — ce qui décroît par construction. Sur les données "
        "réelles : 225 à l'entrée, 32 signées, et un taux de passage jamais supérieur à "
        "100 %. Le dernier, gagnée → signée à 36 %, est en soi une information métier.")
    add_cue(doc,
        "Deux statuts ont été ajoutés après vérification auprès du métier : « Offre "
        "signée », qui n'était dans aucune étape alors que c'est l'étape finale (32 "
        "opportunités, 13,1 M DT, invisibles), et « En attente du plan de charge », "
        "présent dans les données mais absent de la liste blanche — donc rejeté à "
        "chaque chargement. Ce genre de question ne se devine pas : elle se pose.")

    add_h2(doc, "9.3 Le type de graphique est arbitré sur les données")
    add_body(doc,
        "Le modèle propose une forme à partir des mots de la question ; choose_chart_type "
        "la confronte ensuite à la forme réelle des données. « Répartition par pays » "
        "contient bien le mot qui appelle un camembert, mais il y a 19 pays : au-delà de "
        "six parts, deux angles voisins ne se comparent plus. La cardinalité est mesurée "
        "APRÈS application des filtres — une dimension à 19 valeurs peut n'en avoir que "
        "trois sur un périmètre restreint, et le camembert redevient alors le bon choix.")
    add_table(doc, ["Situation", "Forme retenue", "Pourquoi"], [
        ["Camembert sur plus de 6 valeurs", "Barres", "Les angles ne se distinguent plus."],
        ["Camembert d'une moyenne", "Barres",
         "Une moyenne ne s'additionne pas : elle ne forme pas des parts d'un tout."],
        ["Camembert d'un top N", "Barres",
         "Les parts ne totalisent pas le portefeuille, les pourcentages seraient faux."],
        ["Courbe sur une dimension non temporelle", "Barres",
         "Relier des catégories suggère une progression qui n'existe pas."],
        ["Barres sur une dimension temporelle", "Courbe",
         "La question porte sur le mouvement d'ensemble, pas sur des comparaisons deux à deux."],
        ["Tableau, entonnoir, nuage, carte de chaleur", "Inchangée",
         "Formes demandées nommément : les remplacer trahirait la question."],
    ])
    add_body(doc,
        "Quand la forme change, la raison est affichée sous le graphique. Voir des "
        "barres après avoir demandé un camembert, sans explication, donne le sentiment "
        "que l'outil n'a pas compris la question.")

    add_h2(doc, "9.4 Les retouches sont appliquées par le code, pas par le modèle")
    add_body(doc,
        "« En camembert », « top 5 », « par practice », « sans filtre » ne sont pas des "
        "questions : ce sont des ajustements du tableau de bord affiché. Les traiter "
        "comme des questions neuves faisait perdre le contexte — « par pays » repartait "
        "sur le nombre d'opportunités alors que la question d'avant portait sur le "
        "budget — et consommait une requête du quota gratuit pour un changement que le "
        "code sait appliquer instantanément.")
    add_body(doc,
        "Ces retouches se TAPENT, elles ne se cliquent pas. Une rangée de raccourcis "
        "avait été essayée puis retirée : décrire ce que l'on veut est plus expressif "
        "qu'un choix parmi quatre, et la présence des boutons suggérait à tort que le "
        "reste ne se demandait pas.")
    add_cue(doc,
        "Le garde-fou est le cœur du mécanisme. Une retouche hérite EN SILENCE de tout "
        "le contexte précédent : si un seul mot de la demande n'a pas été compris, "
        "hériter revient à répondre à une autre question. Chaque mot porteur de sens "
        "doit donc avoir été consommé par une règle ; au moindre reliquat, la main est "
        "rendue au chemin complet. Le premier jet ne vérifiait que la longueur, et "
        "servait « Risk Advisory » à quelqu'un qui demandait le Maroc.")
    add_body(doc,
        "Une règle secondaire mérite d'être notée : changer d'axe retire le filtre à "
        "valeur unique qui figeait cette même colonne, sinon grouper PAR practice alors "
        "qu'un filtre fige UNE practice donnerait un graphique à une seule barre. Les "
        "filtres à plusieurs valeurs sont épargnés — « compare la France et le Maroc » "
        "veut précisément cet axe ET cette restriction.")

    add_h2(doc, "9.5 « Offre remise » est un terme métier, pas un statut")
    add_body(doc,
        "Le statut décrit l'état COURANT d'une opportunité, pas son historique. Une "
        "offre partie chez le client et gagnée depuis n'est plus au statut « Offre "
        "remise » : compter ce seul statut donne 4 offres là où 57 ont réellement été "
        "déposées. Sont donc « remises » les opportunités dont le statut ATTESTE le "
        "dépôt — Offre remise, En attente du plan de charge, Offre gagnée, Offre "
        "signée, Offre perdue.")
    add_cue(doc,
        "Trois définitions étaient défendables selon qu'on y range les appels d'offres "
        "infructueux et les candidatures non shortlistées — 57, 62 ou 67 offres, pour "
        "un taux de réussite de 53 %, 48 % ou 45 %. L'écart est trop grand pour être "
        "tranché par le code : la question a été posée au métier, qui a retenu la "
        "définition la plus stricte. Il n'existe par ailleurs aucune colonne « date de "
        "remise » : l'échéance en tient lieu, la date limite de dépôt étant, pour un "
        "appel d'offres, la date de remise.")
    add_body(doc,
        "Filtrer sur ces statuts lève au passage l'exclusion par défaut des affaires "
        "perdues (section 9.1) — ce qui est exactement voulu : une offre perdue a bien "
        "été remise, et sans cela la question « sur le total remis, combien de "
        "perdues ? » n'aurait pas de réponse. La définition sert le dashboard ET le "
        "chat : sans règle dédiée, la même question posée dans la conversation aurait "
        "répondu 4.")
    add_body(doc,
        "Le taux de réussite se calcule sur les offres DÉCIDÉES, pas sur le total "
        "remis : une offre en attente n'est ni un succès ni un échec, la compter au "
        "dénominateur écraserait le taux sans rien dire de vrai.")

    add_h2(doc, "9.6 Affaires chaudes, et le piège du range_filter")
    add_body(doc,
        "Une « affaire chaude » est une affaire qu'on va PROBABLEMENT gagner et qui "
        "n'est PAS ENCORE gagnée : statut « Offre remise », OU probabilité de gain "
        "dans [80 %, 100 %[. Une réunion, pas une intersection — l'un des deux "
        "critères suffit. La borne haute est exclue parce que 100 % n'est pas une "
        "prévision dans ces données mais un constat : les 88 lignes à 1,0 sont toutes "
        "« Offre gagnée » ou « Offre signée ». Le portefeuille chaud compte ainsi 9 "
        "opportunités — celles qu'il reste à aller chercher. Les offres perdues en "
        "sortent d'elles-mêmes sans qu'aucun filtre soit nécessaire : ni leur statut "
        "ni leur pondération, vide dans le Sheet, ne les retiennent.")
    add_cue(doc,
        "Ce périmètre a d'abord été restreint aux offres encore en jeu (14 affaires), "
        "puis élargi sur décision métier. La question qui l'a déclenché — « pourquoi "
        "des offres à 100 % n'apparaissent-elles pas ? » — a d'ailleurs mis au jour "
        "une convention de saisie : la pondération est une PRÉVISION tant que l'offre "
        "est ouverte (40 % ou 80 %), un CONSTAT à 100 % une fois gagnée, et vide une "
        "fois perdue. Les 88 opportunités à 100 % sont exactement les 88 déjà gagnées "
        "ou signées, sans une exception.")
    add_body(doc,
        "L'indicateur ne mesure donc plus un pipeline à pousser mais un portefeuille "
        "à forte confiance, mêlant l'acquis et l'à-venir. Les intitulés le disent : "
        "« Budget à forte confiance » plutôt que « Budget en jeu », qui laissait "
        "entendre une affaire encore ouverte.")
    add_body(doc,
        "Le détail occupe sa propre ligne, répartie en TROIS tables côte à côte. Le "
        "tableau de DAC ne défilant pas verticalement, sa hauteur est son nombre de "
        "lignes : la répartir en trois la divise par trois sans en retirer une seule, "
        "ce qui était la contrainte. NTILE(3) plutôt qu'un seuil de rang en dur, le "
        "total changeant avec le filtre de période — 35/35/35 sur l'historique, "
        "19/18/18 par défaut, rangs continus et sans doublon. Le rang est affiché : "
        "trois listes triées côte à côte ne diraient pas, sinon, laquelle vient en "
        "premier.")
    add_cue(doc,
        "Le widget tableau de Bruin DAC ne défile pas verticalement : son élément "
        "externe est un simple overflow-x-auto, sans hauteur ni overflow-y, dans un "
        "cadre overflow-hidden. Borner la hauteur de la ligne le CLIPPE au lieu de le "
        "rendre défilable, et aucun contournement n'existe — le schéma refuse height "
        "sur un widget, le thème refuse le CSS, le TSX partage le même modèle que le "
        "YAML, et le markdown d'un widget text est rendu sans rehype-raw, donc sans "
        "HTML brut. Un rendu React hors de l'iframe a été essayé puis retiré : il "
        "donnait la molette mais sortait le tableau du dashboard.")
    add_cue(doc,
        "Les montants sont en dinars, et DAC ne sait pas l'ecrire a cote du nombre : "
        "son schema refuse suffix, unit et prefix sur une valeur, et son format "
        "currency applique le prefixe $ de son formateur d3 — soit des dollars "
        "affiches la ou il s'agit de dinars. L'unite vit donc sur les LIBELLES (nom "
        "de KPI, titre d'axe, en-tete de colonne), ou elle se lit tout aussi bien "
        "sans empecher le nombre de rester triable et aligne a droite.")
    add_body(doc,
        "Le tableau passant par pandas et les KPI par du SQL, un test confronte les "
        "deux définitions sur les mêmes lignes. Il a révélé un piège de DuckDB : "
        "« NaN >= 0.8 » y vaut VRAI. Le vrai export écrit des NULL, donc rien n'est "
        "faux en production, mais s'il écrivait des NaN le KPI des affaires chaudes "
        "compterait les 170 lignes sans pondération "
        "sans qu'aucune requête n'échoue. Un test l'interdit désormais.")
    add_body(doc,
        "L'application connaissait déjà le terme « offre pondérée » avec la même "
        "intention mais un périmètre plus étroit. Deux définitions divergentes pour "
        "une même réalité auraient donné deux chiffres selon le mot employé : elles "
        "sont unifiées dans business_rules.HOT_DEAL_STATUSES, partagée par le "
        "dashboard et par le chat.")
    add_cue(doc,
        "Brancher ce terme a révélé un même défaut logé dans TROIS modules. "
        "db_layer, sql_builder et response_builder ne s'accordaient pas sur ce que "
        "signifie un range_filter : deux le traitaient comme une demande de liste "
        "brute, le troisième non. « Affaires chaudes par practice » renvoyait donc "
        "les lignes brutes dans le chat pendant que le dashboard groupait, et "
        "« combien d'affaires chaudes ? » répondait « 1 opportunité » — le nombre de "
        "LIGNES du résultat agrégé, au lieu des 14 qu'il contenait.")
    add_body(doc,
        "Le seuil est un MINIMUM : 90 % ou 100 % entrent aussi dans les affaires "
        "chaudes. Encore fallait-il le prouver — sur les vraies données, aucune "
        "affaire encore en jeu ne dépasse 80 %, donc remplacer « >= » par « = » "
        "n'aurait changé aucun chiffre et serait passé inaperçu. dac check prouve "
        "qu'une requête s'exécute, pas qu'elle sélectionne les bonnes lignes : un "
        "fichier de tests dédié exécute désormais le SQL de la vue d'ensemble sur des "
        "lignes fabriquées, et la protection a été vérifiée par mutation.")
    add_cue(doc,
        "Découverte au passage : dans ce Sheet, 100 % n'est pas une prévision mais un "
        "CONSTAT. Les 88 opportunités à 1,0 sont exactement les 88 déjà gagnées ou "
        "signées — correspondance parfaite. Aucune affaire chaude n'apparaît donc à "
        "100 %, non par un défaut du filtre, mais parce qu'une affaire dont la "
        "décision est tombée n'est plus en jeu.")
    add_body(doc,
        "Borner une valeur n'est pas demander une liste. Le souhait d'une liste "
        "s'exprime par use_raw_table ou chart_type == « table », et les trois modules "
        "appliquent maintenant cette règle. Deux tests de non-régression la "
        "verrouillent, un par module concerné.")

    add_h2(doc, "9.7 Compléter plutôt que remplacer")
    add_body(doc,
        "« Ajoute le budget par pays » place un widget de plus sur le tableau de bord "
        "affiché au lieu de le recomposer. Le verbe ne change rien à l'analyse "
        "demandée — même métrique, même axe, mêmes filtres — seulement la façon dont "
        "le résultat rejoint l'écran : d'où un simple drapeau porté par l'intention, "
        "plutôt qu'un chemin de composition parallèle à maintenir.")
    add_cue(doc,
        "Limite acceptée volontairement : l'ajout ne complète PAS la vue d'ensemble, "
        "seulement le tableau de bord de travail. Les widgets de la vue d'ensemble "
        "suivent les filtres de période et de practice de la page, qu'un widget ajouté "
        "ne connaît pas — déplacer le filtre mettrait 24 widgets à jour et pas le 25e. "
        "Un piège silencieux vaut moins qu'une limite énoncée.")
    add_body(doc,
        "Un ajout déjà présent est signalé comme tel plutôt qu'annoncé à tort, et un "
        "plafond de 24 widgets empêche la page de devenir un mur où l'on ne trouve "
        "plus rien.")

    add_h2(doc, "9.8 Dire ce qui a changé")
    add_body(doc,
        "Une demande de suite modifie le tableau de bord affiché. Sans le dire, "
        "l'utilisateur voit l'iframe se recharger sans savoir ce qui a été pris en "
        "compte. La phrase est construite en comparant les deux intentions — jamais "
        "par le modèle : elle décrit ce qui a réellement changé dans les requêtes, pas "
        "ce qu'on croit avoir compris.")
    add_cue(doc,
        "Deux cas restaient muets, et seule l'exécution les a montrés. « En camembert » "
        "sur dix-neuf pays ne produisait AUCUN message : la forme était bien refusée, "
        "mais la raison ne vivait que dans le tableau de bord, et rien d'autre n'ayant "
        "bougé, la comparaison ne trouvait rien à dire. Une demande sans effet — un "
        "filtre déjà posé — était tout aussi silencieuse. Chaque demande reçoit "
        "désormais une réponse explicite.")
    add_body(doc,
        "Au-delà de trois différences, rien n'est annoncé : le mot « modifié » "
        "deviendrait trompeur, ce n'est plus une retouche mais une autre analyse.")

    add_h2(doc, "9.9 Un seul cadre, cinq sections, modifié sur place")
    add_body(doc,
        "Chaque question réécrit _principal.yml, dont le nom de dashboard ne change "
        "jamais — donc l'URL de l'iframe non plus. L'utilisateur voit son tableau de "
        "bord se transformer au lieu d'en voir apparaître un de plus. En parallèle, un "
        "instantané _analyse_<hash>.yml fige le résultat de cette question précise : "
        "c'est ce qui permet à la liste déroulante du chat de rouvrir une analyse "
        "passée telle qu'elle était, sans repasser par le modèle.")
    add_body(doc,
        "Le nom de l'instantané vient de l'objectif résolu et non de la phrase tapée : "
        "nommer un tableau de bord « En camembert » ne dirait rien de ce qu'il montre, "
        "et deux formulations aboutissant à la même analyse partagent ainsi un seul "
        "fichier au lieu d'en accumuler deux identiques.")
    add_body(doc,
        "Le tableau de bord principal, lui, est découpé en cinq dashboards (accueil.yml "
        "et section_*.yml) que la barre d'onglets fait défiler. Le découpage n'est pas "
        "un choix d'esthétique : DAC ne gère pas les dashboards multi-pages — la clé "
        "« pages » est rejetée par son schéma — et douze rangées à faire défiler en "
        "sachant où regarder valaient moins que cinq pages qu'on choisit. Les cinq "
        "fichiers sont découpés par scripts/generate_accueil.py à partir du même "
        "document assemblé, aux mêmes intertitres qui séparaient déjà les groupes : les "
        "règles métier restent donc écrites en un seul endroit.")

    add_h2(doc, "9.10 « Offres gagnées » est un terme métier, comme « offre remise »")
    add_body(doc,
        "La question « combien d'offres gagnées ? » donnait tantôt 56, tantôt 88 selon "
        "la formulation : le modèle proposait parfois le statut littéral « Offre "
        "gagnée », parfois l'ensemble des issues favorables — qui comprend aussi « Offre "
        "signée ». Le terme est désormais tranché par le code, comme « offre remise » "
        "l'était déjà : il vaut toujours les deux statuts, quelle que soit la phrase.")
    add_body(doc,
        "La correction a mis au jour la même asymétrie sur la négation : « offres non "
        "gagnées » n'écartait qu'un des deux statuts, si bien que les offres signées "
        "comptaient dans les DEUX réponses et que les deux questions inverses ne "
        "partitionnaient pas le portefeuille. Les deux formes sont désormais traitées "
        "ensemble — 88 gagnées et 141 non gagnées redonnent bien les 229 opportunités "
        "actives.")

    add_h2(doc, "9.11 Réutiliser une section : une table prouvée, jamais supposée")
    add_body(doc,
        "Composer un tableau de bord pour une question à laquelle une section répond "
        "déjà fait perdre deux choses : l'utilisateur quitte une page qu'il connaît, et "
        "le chiffre affiché provient d'une composition faite à la volée plutôt que d'un "
        "widget relu. backend/overview_match.py décide donc, avant toute écriture, si "
        "une section porte la réponse — et laquelle.")
    add_body(doc,
        "La première version de cette table associait (métrique, axe) à une page en "
        "supposant qu'une même métrique sur un même axe désignait la même question. "
        "C'est faux : le périmètre diffère. Vérification faite widget par widget, cinq "
        "des huit entrées désignaient une page répondant à autre chose — « budget par "
        "pays » menait à une section sans aucun widget par pays, et « combien d'offres » "
        "à un indicateur valant 147 quand le chat en annonce 229 (les offres remises "
        "dont l'échéance est passée, contre le portefeuille actif). Les deux chiffres "
        "sont justes ; la page contredisait la phrase qu'on venait de lire.")
    add_body(doc,
        "La leçon porte moins sur la table que sur la méthode. Elle ne conserve que les "
        "entrées prouvées, et la preuve est devenue un test : tests/test_reutilisation_"
        "fidele.py exécute les deux moteurs sur les mêmes lignes et refuse toute entrée "
        "dont le total ou le nombre de lignes diffère, avec un contrôle négatif qui "
        "vérifie que le test sait dire non. Ajouter une entrée fausse fait désormais "
        "tomber la suite au lieu d'atteindre l'utilisateur.")

    add_h1(doc, "10. Réponses textuelles déterministes")
    add_body(doc,
        "response_builder.py calcule le message affiché à l'utilisateur directement "
        "depuis les données retournées par la requête SQL — jamais par le LLM. "
        "format_metric_value() applique les règles d'unité (budget en €, "
        "win_probability en % en multipliant par 100 car stocké en base comme "
        "fraction 0-1, nb_opportunities en entier). Cette séparation garantit que le "
        "texte affiché ne peut jamais diverger des chiffres réels du graphique.")

    # --- Alertes ---
    add_h1(doc, "11. Alertes deadlines")
    add_h2(doc, "11.1 Une donnée « jours restants » qui ne peut pas devenir obsolète")
    add_body(doc,
        "Historiquement, days_remaining était une colonne figée à l'import, qui ne "
        "diminuait pas toute seule au fil des jours — il fallait donc un job nocturne "
        "pour la recalculer, et se méfier de sa valeur entre deux exécutions. Depuis le "
        "passage à une lecture directe du Sheet, le problème disparaît par construction : "
        "days_remaining est recalculée à CHAQUE chargement des données (toutes les 15 "
        "minutes) depuis la deadline réelle et la date du jour. Il n'y a plus rien à "
        "rafraîchir la nuit, et le module de maintenance dédié a été supprimé.")
    add_body(doc,
        "C'est un bon exemple de simplification obtenue en changeant la structure "
        "plutôt qu'en ajoutant du code : le mécanisme le plus fiable est ici celui "
        "qu'on a pu retirer.")
    add_h2(doc, "11.2 Un job planifié (APScheduler, backend/main.py)")
    add_table(doc,
        ["Fréquence", "Job", "Rôle"],
        [
            ["Toutes les 15 min", "refresh_dataframe", "Recharge les données depuis le Sheet et "
             "régénère la projection DuckDB. Recalcule au passage days_remaining."],
            ["08:00", "run_daily_alert_check_if_needed", "Sélectionne les opportunités actives à "
             "échéance ≤ 7 jours et envoie l'email récapitulatif si la liste n'est pas vide."],
        ])
    add_body(doc,
        "Les deux tournent aussi une fois au démarrage, en tâche de fond : le serveur "
        "accepte les requêtes immédiatement, sans attendre le premier chargement du "
        "Sheet (qui prend une à deux secondes).")
    add_h2(doc, "11.4 Envoi de l'email (Gmail SMTP)")
    add_body(doc,
        "send_alert_email() utilise smtplib avec STARTTLS sur smtp.gmail.com:587, "
        "authentifié par un mot de passe d'application Gmail (jamais le mot de passe "
        "principal du compte) — trois variables d'environnement dédiées "
        "(GMAIL_SENDER, GMAIL_APP_PASSWORD, ALERT_RECIPIENT_EMAIL), absentes du dépôt "
        "git. Le corps de l'email liste chaque opportunité avec son urgence, son "
        "client, son statut et son budget.")
    add_h2(doc, "11.5 Exclusion des statuts clos")
    add_body(doc,
        "EXCLUDED_STATUSES filtre les opportunités déjà gagnées, perdues, signées, "
        "infructueuses, « NO GO », hors scope ou non shortlistées — une opportunité "
        "close ne doit plus jamais déclencher d'alerte, même si sa date d'échéance "
        "reste techniquement proche.")
    add_h2(doc, "11.6 Bandeau frontend (AlertBanner.jsx)")
    add_body(doc,
        "Le composant interroge GET /alerts/deadlines au montage (lecture live, "
        "indépendante du digest email quotidien — mêmes règles métier) et affiche un "
        "bandeau dépliable en haut du dashboard, avec un code couleur d'urgence "
        "(rouge ≤ 3 jours, orange 4-7 jours) réutilisant les jetons de statut déjà "
        "définis dans le design system de l'application.")

    add_h2(doc, "11.7 Bug trouvé : une échéance déjà passée dans une liste « urgente »")
    add_body(doc,
        "En testant la question « liste des opportunités urgentes (< 7 jours) », des "
        "opportunités dont l'échéance était dépassée depuis des mois apparaissaient dans "
        "les résultats. Cause : la traduction en filtre SQL était days_remaining < 7, "
        "et une deadline dépassée depuis 50 jours a days_remaining = -50, donc -50 < 7 "
        "est vrai — mathématiquement correct, sémantiquement faux : « urgent » sous-entend "
        "une échéance encore à venir.")
    add_body(doc,
        "Corrigé à trois niveaux plutôt qu'un seul, par prudence — le même principe de "
        "défense en profondeur que le reste de l'anti-hallucination du projet :")
    add_numbered(doc, "le chemin rapide par mots-clés (intent_refiner.py) produit désormais "
                       "{\"op\": \"between\", \"value\": [0, N]} au lieu de {\"op\": \"<\", \"value\": N} ;")
    add_numbered(doc, "l'exemple few-shot et une instruction explicite ont été ajoutés au prompt "
                       "système du LLM (llm.py) pour le même cas ;")
    add_numbered(doc, "un garde-fou déterministe dans refine_intent() normalise automatiquement "
                       "tout days_remaining avec un opérateur \"<\" ou \"<=\" en \"between\" à partir "
                       "de 0 — quelle que soit son origine (LLM ou règle), et même si les deux "
                       "premiers points étaient un jour contournés.")
    add_body(doc,
        "Le troisième niveau est le plus important pédagogiquement : une instruction de prompt "
        "n'est jamais une garantie (le LLM l'a déjà ignorée une fois pour le scatter, voir "
        "section 15), donc la correction réellement fiable est celle qui ne dépend pas du "
        "LLM pour se déclencher. Le tableau affiche aussi désormais « Échéance dépassée "
        "depuis N jours » plutôt qu'un nombre négatif brut, pour les cas où une deadline "
        "passée reste légitimement visible (ex: une liste non filtrée par date).")

    add_h2(doc, "11.8 Statuts clos exclus des requêtes « urgentes » du chat")
    add_body(doc,
        "Signalé par l'utilisateur : « liste des opportunités urgentes » pouvait afficher "
        "une offre déjà gagnée ou perdue si sa date d'échéance technique tombait dans la "
        "fenêtre demandée — incohérent avec les alertes email/bannière (section 10.5), qui "
        "excluent bien ces statuts. Plutôt que de dupliquer la liste EXCLUDED_STATUSES, "
        "intent_refiner.py::refine_intent() l'importe directement depuis alerts.py et pose "
        "intent[\"exclude_statuses\"] dès qu'un filtre days_remaining est présent — jamais "
        "piloté par le LLM, uniquement par ce garde-fou déterministe. db_layer.py traduit "
        "ce champ en une clause status NOT IN (...), ajoutée à toute requête qui filtre sur "
        "l'urgence. Une seule définition de « actif » pour tout le projet, jamais deux "
        "listes de statuts qui pourraient diverger avec le temps.")

    add_h2(doc, "11.9 Rattrapage du scheduler")
    add_body(doc,
        "APScheduler ne rattrape pas un job manqué si le processus était entièrement arrêté "
        "au moment prévu (contrairement à misfire_grace_time, qui ne couvre qu'un léger "
        "retard sur un processus resté actif). Pour le rafraîchissement des données, ce "
        "n'est pas un problème : il est rappelé sans condition à chaque démarrage, et le "
        "refaire ne coûte rien. Pour l'email quotidien, un rappel sans condition renverrait "
        "un doublon si le cron de 8h avait déjà tourné le jour même.")
    add_body(doc,
        "Solution : un fichier `data/scheduler_state.json` retenant la date de dernière "
        "exécution par job — auparavant une table MySQL, devenue un simple fichier local "
        "avec la suppression de la base. run_daily_alert_check_if_needed() vérifie d'abord "
        "si le job a déjà tourné aujourd'hui ; sinon, elle exécute la vérification puis "
        "marque la date. Appelée à la fois par le cron de 8h ET au démarrage du serveur : "
        "un serveur éteint pile à 8h rattrape l'envoi dès qu'il redémarre, quelle que soit "
        "l'heure, sans jamais produire de second email le même jour.")
    add_body(doc,
        "Détail de robustesse : si la vérification anti-doublon elle-même échoue (fichier "
        "illisible), l'alerte est envoyée quand même. Un mécanisme de suivi cassé ne doit "
        "jamais avoir pour conséquence de supprimer silencieusement l'alerte qu'il protège.")

    # --- Pont DuckDB ---
    add_h1(doc, "12. Le pont DuckDB (backend/duckdb_export.py)")
    add_body(doc,
        "Bruin DAC ne sait interroger que des connexions SQL. Les données vivant "
        "désormais dans un DataFrame en mémoire, un pont est nécessaire : ce module "
        "réécrit le DataFrame dans un fichier DuckDB local à chaque rafraîchissement.")

    add_h2(doc, "12.1 Une projection, pas une base de données")
    add_body(doc,
        "La distinction est importante : le fichier DuckDB n'est jamais lu par le code "
        "applicatif, jamais édité à la main, et intégralement régénéré à chaque cycle. "
        "pandas reste la source de vérité pour le chat et les alertes. Cette asymétrie "
        "garantit qu'aucune divergence ne peut s'installer entre ce que dit le chat et "
        "ce qu'affichent les tableaux de bord — les deux dérivent du même chargement.")
    add_body(doc,
        "L'écriture utilise CREATE OR REPLACE TABLE plutôt qu'un DELETE suivi d'INSERT : "
        "atomique du point de vue d'un lecteur, et le schéma est repris automatiquement "
        "du DataFrame — aucune définition de colonnes à maintenir en double.")

    add_h2(doc, "12.2 Le piège de concurrence, identifié puis testé")
    add_body(doc,
        "DuckDB autorise SOIT plusieurs lecteurs, SOIT un seul écrivain — jamais les "
        "deux simultanément. Or DAC lit le fichier pendant que l'application le "
        "réécrit toutes les 15 minutes : une collision était donc possible par "
        "construction.")
    add_body(doc,
        "Deux mesures. Côté DAC, la connexion est déclarée en lecture seule "
        "(read_only: true dans .bruin.yml) — sans quoi elle prendrait un verrou "
        "exclusif et empêcherait tout rafraîchissement. Côté écriture, la connexion est "
        "fermée immédiatement après usage, et l'opération réessaie quelques fois à "
        "intervalle court si le fichier est momentanément verrouillé. En cas d'échec "
        "complet, l'ancien fichier est conservé : les tableaux de bord affichent des "
        "données d'un cycle plus anciennes, jamais une base vide ou corrompue.")
    add_body(doc,
        "Vérifié sous charge plutôt que supposé : trois rafraîchissements concurrents "
        "lancés pendant vingt-sept requêtes de widgets, tous passants.")

    add_h2(doc, "12.3 Authentification Google — compte de service, pas clé API")
    add_body(doc,
        "L'application doit pouvoir ÉCRIRE dans le Sheet (identifiants des nouvelles "
        "lignes et colonnes calculées, section 4.3), ce qu'une simple clé API ne permet "
        "pas — elle est en lecture seule. L'accès passe donc par un compte de service "
        "Google (fichier JSON, ajouté en Éditeur sur le Sheet cible).")
    add_body(doc,
        "Optimisation mesurée : la résolution du Sheet (open_by_key puis worksheet) "
        "coûte deux allers-retours réseau, soit 1,6 s, et était rejouée à chaque cycle "
        "alors que le client authentifié, lui, était déjà mis en cache. Mettre aussi en "
        "cache l'objet Worksheet fait tomber ce coût à zéro sur les appels suivants — "
        "un chargement complet passe de ~1,1 s à ~0,4 s. Le profilage avait d'ailleurs "
        "démenti l'intuition de départ : les écritures en base ne représentaient que "
        "0,03 s, l'essentiel du temps était réseau.")

    add_h2(doc, "12.4 Déclenchement")
    add_body(doc,
        "Trois déclencheurs pour le cycle complet (lecture du Sheet, reconstruction du "
        "DataFrame, projection DuckDB, purge du cache de tableaux de bord) : toutes les "
        "15 minutes via APScheduler, une fois au démarrage, et à la demande via "
        "POST /sheets/sync — ce dernier étant exposé par un bouton dans l'interface de "
        "chat, qui affiche le résumé du chargement (lignes chargées, ignorées).")

    # --- Frontend ---
    add_h1(doc, "13. Frontend (Vite + React)")
    add_bullet(doc, "Composants principaux : ChatPanel, DashboardPanel (iframe DAC), AlertBanner, "
                     "DataTable, StatTile, ThemeToggle, DevoteamLogo.")
    add_bullet(doc, "Hooks dédiés : useTheme (clair/sombre), useChatHistory (persistance "
                     "localStorage).")
    add_bullet(doc, "DashboardPanel affiche un iframe pointant sur le serveur DAC : la vue "
                     "d'ensemble par défaut, ou le tableau de bord généré dès qu'une question a été "
                     "posée. Deux boutons permettent de basculer entre les deux. L'URL est "
                     "construite à partir du NOM du tableau de bord (DAC route par nom affiché, pas "
                     "par nom de fichier), et le compteur dashboardKey sert de clé React pour forcer "
                     "le rechargement de l'iframe même quand le nom ne change pas.")
    add_bullet(doc, "Thème géré par variables CSS (custom properties) — le mode sombre est "
                     "SÉLECTIONNÉ avec ses propres valeurs, pas un simple filtre inversé "
                     "automatique. Il s'applique au chat, pas à l'iframe DAC (voir section 17).")
    add_bullet(doc, "Jetons de marque (--brand-primary…) volontairement séparés des jetons "
                     "de sécurité des graphiques (--series-1…8) pour ne jamais mélanger "
                     "esthétique et lisibilité daltonisme.")
    add_bullet(doc, "Astuce React : un compteur dashboardKey, incrémenté à chaque nouvelle "
                     "réponse, est utilisé comme prop key pour forcer le remontage des "
                     "composants et rejouer les animations d'entrée à chaque nouveau résultat.")

    add_h2(doc, "13.1 Historique de conversation persistant (useChatHistory.js)")
    add_body(doc,
        "Messages, dashboard affiché et contexte multi-tour (previous_intent) sont "
        "sérialisés dans localStorage à chaque changement, et restaurés au montage — "
        "plafonnés à 100 messages pour ne pas faire grossir indéfiniment le stockage. Le "
        "compteur d'identifiants de message (nextIdRef) est réinitialisé au-delà du plus "
        "grand id restauré, pour ne jamais entrer en collision avec l'historique repris. "
        "Toute erreur (quota dépassé, navigation privée, JSON corrompu) est absorbée "
        "silencieusement : la persistance est un confort, jamais une dépendance dont "
        "l'absence casserait le chat. Un bouton « Effacer la conversation » (avec "
        "confirmation) a été ajouté dans le header, puisqu'un historique qui persiste "
        "indéfiniment a besoin d'une porte de sortie explicite.")

    add_h2(doc, "13.2 La saisie en tête de panneau et le volet d'historique")
    add_body(doc,
        "La zone de saisie était en pied de panneau, sous les messages : elle descendait "
        "avec la conversation et demandait de faire défiler pour être retrouvée. Elle est "
        "désormais en tête, juste sous l'en-tête, et ne bouge plus. Le trait de séparation "
        "passe simplement de border-top à border-bottom.")
    add_body(doc,
        "C'est le SEUL moyen de modifier le tableau de bord affiché. Une rangée de "
        "retouches en un clic avait été essayée puis retirée : décrire ce que l'on veut "
        "est plus expressif qu'un choix parmi quatre raccourcis, et la présence des "
        "boutons suggérait à tort que le reste ne se demandait pas.")
    add_body(doc,
        "Les analyses déjà produites vivent dans un volet d'historique, ouvert depuis "
        "l'en-tête du TABLEAU DE BORD — pas du chat : c'est par lui qu'on choisit "
        "quelle analyse regarder, il précède donc le sélecteur de vue, qui ne fait "
        "ensuite que basculer entre elle et la vue d'ensemble. Superposé au chat, "
        "groupé par jour, il signale l'analyse "
        "affichée et rouvre n'importe quelle demande passée d'un clic. Une liste "
        "déroulante avait d'abord été livrée, puis remplacée — elle cachait son contenu "
        "tant qu'on ne l'ouvrait pas, tronquait des intitulés qui sont des phrases "
        "entières, et ne pouvait porter ni date ni état courant.")
    add_cue(doc,
        "L'historique retrace ce que l'utilisateur a écrit, pas ce que le backend a produit : "
        "deux formulations aboutissant au même tableau de bord y restent deux entrées, "
        "parce que c'est par sa propre phrase qu'on retrouve une analyse. D'où la clé sur "
        "l'identifiant du message et non sur le nom du tableau de bord, qui peut se "
        "répéter. Les conversations enregistrées avant l'ajout de l'horodatage sont "
        "regroupées sous « Plus ancien » plutôt que de porter une fausse date.")

    add_h2(doc, "13.3 Un seul tableau de bord, et des transitions en fondu")
    add_body(doc,
        "La bascule entre « Vue d'ensemble » et « Mon tableau de bord » a été retirée : "
        "il n'y a plus qu'un tableau de bord, qui s'ouvre sur la vue d'ensemble et que "
        "les questions transforment. Le rechargement de la page y revient — "
        "l'affichage et le contexte de la dernière question ne sont plus restaurés, "
        "car les conserver rouvrait une analyse que personne n'avait redemandée et "
        "laissait une suite comme « en camembert » ajuster un tableau de bord qu'on ne "
        "regardait plus. La conversation, elle, reste conservée.")
    add_cue(doc,
        "Recharger l'iframe en place la vidait le temps du chargement : l'écran "
        "passait au blanc à chaque question. Deux cadres se superposent désormais — "
        "celui qu'on regarde, et le suivant qui charge par-dessus, invisible. Quand il "
        "est prêt il devient opaque, et l'ancien n'est retiré qu'ENSUITE, une fois le "
        "fondu terminé : retirer l'ancien au moment de la promotion aurait laissé "
        "apparaître le fond entre les deux.")
    add_body(doc,
        "Le squelette animé ne sert donc plus qu'au tout premier affichage, quand il "
        "n'y a effectivement rien à montrer en attendant (démarrage à froid du moteur "
        "de requête). Les fois suivantes, le tableau de bord précédent tient ce rôle "
        "bien mieux, et un fil d'attente de deux pixels en tête signale le travail en "
        "cours. La préférence système « mouvement réduit » ramène le tout à un "
        "basculement instantané, sans jamais rouvrir de fenêtre de vide.")

    add_h2(doc, "13.4 Bug trouvé : header illisible en mode sombre")
    add_body(doc,
        "Trouvé en relisant le CSS (pas en testant à l'œil — aucun outil de capture "
        "d'écran disponible dans cet environnement) : le dégradé du header utilisait "
        "linear-gradient(135deg, var(--brand-primary) 0%, var(--text-primary) 68%). "
        "--text-primary vaut #0b0b0b (presque noir) en mode clair, mais #ffffff (blanc) "
        "en mode sombre — et le texte du header, lui, est en blanc fixe (color: #ffffff). "
        "En mode sombre, le dégradé finissait donc en blanc, sous un texte blanc : "
        "illisible, tout comme le logo (lui aussi dessiné en blanc).")
    add_body(doc,
        "Corrigé avec un nouveau jeton --header-ink, fixé à #0b0b0b et jamais réactif au "
        "thème — le header porte toujours du texte blanc, donc la fin de son dégradé doit "
        "toujours rester sombre, quel que soit le thème actif. Un second bug de même "
        "nature affectait le texte « warning » du bandeau d'alerte "
        "(color-mix(..., black) littéral, qui suppose un fond clair) : corrigé avec un "
        "jeton --warning-ink-mix (black en mode clair, white en mode sombre). Aucune "
        "valeur n'a changé en mode clair dans les deux cas — uniquement des ajouts, "
        "jamais une modification du rendu déjà validé.")

    # --- Sécurité ---
    add_h1(doc, "14. Sécurité")
    add_bullet(doc, "Aucune injection SQL possible : le LLM ne produit jamais de SQL, "
                     "uniquement un JSON borné par une liste blanche, converti en requêtes "
                     "paramétrées.")
    add_bullet(doc, ".env (clés API, identifiants email) exclu du suivi git via .gitignore ; "
                     ".env.example documente les variables attendues sans valeurs réelles.")
    add_bullet(doc, "Le mot de passe d'application Gmail est distinct du mot de passe principal "
                     "du compte et révocable indépendamment.")
    add_bullet(doc, "credentials/ (clé de compte de service Google, section 11.1) est gitignoré "
                     "au niveau du dossier entier — jamais de fichier JSON de credentials commité, "
                     "quel que soit son nom.")

    # --- Tests ---
    add_h1(doc, "15. Tests automatisés")
    add_body(doc,
        "337 tests pytest, sans dépendance réseau ni données réelles : le client Gemini "
        "et le client Google Sheets (gspread) sont simulés (monkeypatch), et les données "
        "sont un petit DataFrame construit dans le test. La suite tourne donc hors ligne, "
        "en quelques secondes.")
    add_body(doc,
        "Ces tests sont systématiquement complétés par une vérification en conditions "
        "réelles — Sheet réel, appel Gemini réel, exécution effective de chaque widget "
        "via `dac check`. La répartition des rôles est nette : les tests protègent contre "
        "les régressions, la vérification réelle est ce qui a révélé la quasi-totalité "
        "des vrais bugs de ce projet (sections 8.3 et 15).")
    add_table(doc,
        ["Fichier", "Tests", "Ce qu'il protège"],
        [
            ["test_intent_refiner.py", "62", "Dates relatives, parseur par mots-clés, affinage "
             "d'intention, normalisation funnel/heatmap/scatter/days_remaining/exclude_statuses, "
             "arbitrage du type de graphique, retouches et vérification de couverture mot à mot."],
            ["test_data_store.py", "29", "Parsing et validation par ligne, littéraux NULL, "
             "attribution et réécriture des identifiants, réparation cellule par cellule sans "
             "perdre la ligne, traçabilité de chaque remplacement, résilience à une panne de "
             "lecture du Sheet."],
            ["test_dac_composer.py", "40", "Génération SQL (échappement des apostrophes, refus "
             "d'injection), composition multi-widgets, filtres propagés à tous les widgets, grille "
             "12 colonnes respectée, exclusion des affaires perdues, nom du dashboard de travail "
             "constant, YAML valide et relisible."],
            ["test_db_layer.py", "22", "Requêtage pandas : groupements, filtres à valeurs multiples, "
             "intervalles, exclusion des affaires perdues et son échappatoire, chemin brut du "
             "scatter, DataFrame vide."],
            ["test_llm_validation.py", "16", "Anti-hallucination : clarification plutôt que valeur "
             "devinée, résolution fuzzy des filtres, chemin rapide vs LLM."],
            ["test_response_builder.py", "26", "Formatage des unités, messages texte déterministes, "
             "cohérence texte/graphique pour funnel et heatmap, description de ce qui a changé "
             "dans le tableau de bord."],
            ["test_hot_deals.py", "9", "Le critere des affaires chaudes de bout en bout : seuil inclusif, statut sans effet, ponderation absente ecartee, accord entre la definition pandas et le SQL des KPI, et NULL plutot que NaN a l export."],
            ["test_accueil_dashboard.py", "11", "SQL de la vue d'ensemble exécuté sur des lignes "
             "fabriquées : seuil des affaires chaudes inclusif, issues qui totalisent les offres "
             "remises, taux de réussite calculé sur les seules offres décidées."],
            ["test_palette.py", "3", "La palette de graphiques vit dans deux fichiers qui doivent rester identiques ; chaque teinte atteint 3:1 avec la surface ; aucune n est employee deux fois."],
            ["test_frontend_proxy.py", "3", "Les appels de src/api.js, les prefixes du proxy Vite et les routes de main.py doivent rester d accord ; le montage statique reste declare en dernier."],
            ["test_alerts.py", "12", "Fenêtre de 7 jours, exclusion des statuts clos, envoi email, "
             "rattrapage idempotent du scheduler."],
        ])

    # --- Q&A entretien ---
    add_h1(doc, "16. Questions d'entretien probables")
    qa = [
        ("Pourquoi ne pas laisser le LLM écrire du SQL directement ?",
         "Parce qu'un LLM peut halluciner une colonne, une table ou une valeur inexistante. "
         "En le limitant à produire un JSON borné par une liste blanche stricte "
         "(schema_and_whitelist.py), toute requête SQL exécutée est garantie syntaxiquement "
         "et sémantiquement valide, quelle que soit la sortie du modèle."),
        ("Comment évitez-vous les hallucinations, concrètement ?",
         "Trois filets superposés : (1) Pydantic rejette toute clé de filtre hors liste "
         "blanche, (2) _fuzzy_match ne retourne jamais une valeur absente des données "
         "réelles — sinon IntentUnclear déclenche une clarification, (3) les dates relatives "
         "sont calculées par arithmétique Python déterministe, jamais devinées par le LLM."),
        ("Pourquoi avoir supprimé la base de données ?",
         "Parce qu'elle ne servait plus à rien. Depuis que les équipes saisissent dans le "
         "Google Sheet, MySQL n'était qu'un intermédiaire : on y recopiait le Sheet toutes "
         "les 15 minutes pour le relire juste après. Avec ~360 lignes et 18 colonnes, les "
         "données tiennent en mémoire dans un DataFrame pandas. Bénéfice mesurable au-delà "
         "de l'installation simplifiée : toute la machinerie de vues pré-agrégées (choix de "
         "vue, vérification de compatibilité, repli sur un calcul brut) a disparu, car "
         "pandas groupe uniformément. Ce choix serait à revoir à partir de quelques dizaines "
         "de milliers de lignes, ou s'il fallait des écritures concurrentes."),
        ("Si les données sont dans pandas, pourquoi un fichier DuckDB ?",
         "Parce que le moteur de tableaux de bord (Bruin DAC) ne sait interroger que des "
         "connexions SQL. DuckDB est une projection en lecture seule du DataFrame, "
         "régénérée à chaque rafraîchissement — pas un retour à une base de données : "
         "aucun code applicatif ne la lit. pandas reste la source de vérité, ce qui "
         "garantit que le chat et les tableaux de bord ne peuvent pas afficher des "
         "chiffres différents."),
        ("Pourquoi ne pas laisser l'IA écrire les tableaux de bord librement ?",
         "Elle les écrit — mais pas le SQL. Le modèle produit une intention validée contre "
         "la liste blanche, et du code déterministe la traduit en requêtes. Laisser le "
         "modèle produire le SQL rouvrirait exactement la classe de bugs que tout le projet "
         "empêche : colonne inventée, filtre halluciné, agrégation silencieusement fausse. "
         "Le coût de cette contrainte est nul en pratique, puisque les 18 colonnes fixes du "
         "Sheet sont déjà entièrement décrites par la liste blanche."),
        ("Comment gérez-vous le contexte multi-tour sans session serveur ?",
         "Le frontend renvoie le dernier intent résolu (previous_intent) à chaque nouvelle "
         "question. Dès qu'un contexte est présent, le parseur rapide est systématiquement "
         "court-circuité au profit du LLM, à qui ce contexte est injecté dans le prompt "
         "système pour qu'il décide d'hériter ou non des paramètres précédents."),
        ("Comment garantissez-vous que « jours restants » n'est jamais obsolète ?",
         "Par construction, plus par un job de maintenance. Cette valeur est recalculée à "
         "chaque chargement des données (toutes les 15 minutes) depuis la deadline réelle "
         "et la date du jour. Auparavant, c'était une colonne figée qu'un job nocturne "
         "rafraîchissait — avec le risque classique d'une dérive si le job était manqué. "
         "Le module de maintenance dédié a pu être supprimé : le mécanisme le plus fiable "
         "est ici celui qu'on a retiré."),
        ("Que se passe-t-il si le LLM renvoie un JSON malformé ou un schéma invalide ?",
         "json.JSONDecodeError et pydantic.ValidationError sont interceptées explicitement "
         "et transformées en message de clarification pour l'utilisateur — jamais une "
         "exception non gérée ni une réponse partiellement remplie."),
        ("Comment le système empêche-t-il un graphique illisible (trop de catégories) ?",
         "MAX_PIE_SLICES (6) et MAX_BAR_CATEGORIES (12) : au-delà, un camembert bascule "
         "en barres horizontales, et la traîne est regroupée dans « Autres » pour les "
         "agrégations sum/count (jamais pour une moyenne, qui n'aurait pas de sens agrégée)."),
        ("Une instruction de prompt suffit-elle à garantir un comportement du LLM ?",
         "Non — observé concrètement : une question de corrélation était systématiquement "
         "interprétée comme un simple KPI malgré un exemple few-shot explicite, parce que "
         "le parseur rapide (mots-clés, avant même l'appel LLM) l'interceptait avec une "
         "confiance mal placée. Corrigé en élargissant la détection de mots-clés — la leçon : "
         "diagnostiquer d'abord QUEL composant a décidé avant de corriger le prompt, le "
         "problème n'est pas toujours là où on l'imagine."),
        ("Pourquoi corriger le bug days_remaining à trois endroits différents ?",
         "Défense en profondeur, le même principe que l'anti-hallucination générale du "
         "projet : le chemin rapide et le prompt LLM sont corrigés pour ne plus produire "
         "le bug, ET un garde-fou déterministe dans refine_intent() normalise le résultat "
         "quelle que soit son origine — utile si jamais les deux premiers correctifs sont "
         "un jour contournés par une reformulation imprévue."),
        ("Comment garantis-tu qu'un email d'alerte n'est jamais envoyé deux fois le même jour ?",
         "Un fichier local (data/scheduler_state.json) trace la date de dernière exécution "
         "par job — c'était une table MySQL avant la suppression de la base. "
         "run_daily_alert_check_if_needed() vérifie cette date avant d'agir, et la met à "
         "jour après — appelée à la fois par le cron 8h et au démarrage du serveur, donc "
         "un redémarrage tardif rattrape l'envoi manqué sans risquer un doublon si le cron "
         "avait déjà tourné plus tôt dans la journée. Nuance de robustesse : si la lecture "
         "de ce fichier échoue, l'alerte part quand même — un mécanisme de suivi cassé ne "
         "doit jamais supprimer silencieusement l'alerte qu'il protège."),
        ("Comment as-tu trouvé les bugs de contraste en mode sombre sans navigateur ni capture d'écran ?",
         "En lisant le CSS ligne par ligne plutôt qu'en devinant : chercher chaque endroit "
         "où une couleur qui change avec le thème (var(--text-primary), un color-mix vers "
         "un noir/blanc littéral) est combinée avec une couleur de texte fixe. Le header "
         "avait un texte blanc fixe sur un dégradé qui finissait par --text-primary — "
         "blanc lui aussi en mode sombre. Un raisonnement statique suffit à repérer ce "
         "type de bug ; il n'a fallu aucun rendu réel pour le localiser, seulement pour le "
         "confirmer visuellement ensuite."),
        ("Un test avec des mocks peut-il garantir qu'une intégration externe fonctionne "
         "réellement ?",
         "Non — exemple concret rencontré sur ce projet, à l'époque où les données "
         "transitaient encore par MySQL : tous les tests mockés passaient, mais le premier "
         "essai en conditions réelles rapportait 359 lignes sur 360 comme « id introuvable » "
         "alors qu'elles existaient toutes. La cause (cursor.rowcount qui compte les lignes "
         "modifiées, pas trouvées) n'était visible qu'avec de vraies données où la plupart "
         "des mises à jour ne changent rien. Les mocks avaient simulé le comportement qu'on "
         "ATTENDAIT, pas le vrai — la leçon : les mocks valident la logique connue, seule "
         "une vérification en conditions réelles révèle ce qu'on n'avait pas anticipé."),
        ("Quel bug ce projet vous a-t-il appris à chercher en priorité ?",
         "Ceux qui produisent un résultat plausible plutôt qu'une erreur. Trois exemples "
         "réels : une liste d'opportunités « urgentes » qui incluait des échéances dépassées "
         "(moins sept jours est bien inférieur à sept, mathématiquement juste et "
         "sémantiquement faux) ; un texte qui décrivait dix-neuf statuts quand le graphique "
         "n'en montrait que dix ; et un widget « budget par practice » réduit à une seule "
         "barre parce que la question filtrait déjà sur une practice. Aucun ne lève "
         "d'exception, aucun n'échoue en test unitaire — ils se voient seulement en "
         "regardant le résultat final avec de vraies données."),
        ("Que se passe-t-il si la génération du tableau de bord échoue ?",
         "La réponse textuelle est renvoyée quand même : l'appel est enveloppé dans un "
         "try/except et l'échec est journalisé, pas propagé. Le frontend retombe alors sur "
         "la vue d'ensemble. Même principe pour la projection DuckDB : si le fichier est "
         "momentanément verrouillé, l'ancien est conservé et les tableaux de bord affichent "
         "des données d'un cycle plus anciennes — jamais une erreur à l'écran, jamais une "
         "base vide."),
    ]
    for question, answer in qa:
        add_h3(doc, "Q : " + question)
        add_body(doc, answer)

    # --- Limites ---
    add_h1(doc, "17. Les deux modes d'exécution")
    add_body(doc,
        "L'application se lance de deux façons, et se tromper est silencieux. Un "
        "raccourci de Bureau existe pour chacune (scripts/create_shortcut.ps1 les "
        "crée toutes les deux).")
    add_table(doc,
        ["", "Développement", "Production"],
        [
            ["Lanceur", "scripts/start_dev.bat", "scripts/start_prod.bat"],
            ["Processus", "3 (API, DAC, Vite)", "2 (API + interface, DAC)"],
            ["Frontend", "servi par Vite, port 5173", "compilé, servi par le backend"],
            ["Adresse", "http://localhost:5173", "http://127.0.0.1:8000"],
            ["Rechargement à chaud", "oui (--reload)", "non"],
            ["Workers uvicorn", "1", "1 (imposé)"],
        ])
    add_body(doc,
        "Trois différences méritent leur justification. Le rechargement à chaud "
        "surveille l'arborescence — que l'application modifie elle-même, puisque "
        "chaque question réécrit dac/dashboards/_principal.yml : en production, le "
        "serveur se redémarrerait à chaque réponse qu'il vient de produire. Le "
        "frontend doit être compilé AVANT le démarrage, car le backend sert "
        "frontend/dist tel qu'il le trouve : partir sur une compilation périmée "
        "afficherait l'ancienne interface sans le moindre signe, et le script s'arrête "
        "donc si npm run build échoue. Enfin un seul worker, parce que le jeu de "
        "données vit en mémoire dans le processus et qu'un planificateur y tourne : "
        "plusieurs workers garderaient chacun leur copie des données ET relanceraient "
        "le planificateur — emails d'alerte en double, écritures concurrentes dans le "
        "Google Sheet.")
    add_cue(doc,
        "tests/test_lanceurs.py confronte les deux scripts : ils ne peuvent pas "
        "dériver l'un vers l'autre sans faire échouer les tests.")

    add_h1(doc, "18. Limites connues")
    add_bullet(doc, "Google Gemini (fournisseur actuel depuis la migration forcée, section 2.2) "
                     "supporte en fait les sorties structurées strictes (vérifié empiriquement), "
                     "mais l'app garde volontairement l'architecture JSON libre héritée de Groq "
                     "— le filet de sécurité Pydantic + liste blanche reste donc le mécanisme "
                     "d'application des règles, plutôt qu'une garantie au niveau du schéma du "
                     "modèle. Adopter le mode strict resterait une amélioration possible.")
    add_bullet(doc, "L'historique de conversation persistant (localStorage) est par "
                     "navigateur/appareil, pas partagé entre postes — nécessiterait un compte "
                     "utilisateur pour ça.")
    add_bullet(doc, "Trois processus doivent tourner (API, serveur DAC sur le port 8321, "
                     "frontend) au lieu d'un seul. scripts/start_dev.bat les lance ensemble, mais "
                     "un tableau de bord vide dans l'iframe signifie presque toujours que le "
                     "serveur DAC n'est pas démarré.")
    add_bullet(doc, "Le mode sombre ne s'applique pas à la zone tableau de bord. L'identité "
                     "Devoteam, elle, y est appliquée par un thème dédié (couleurs et palette de "
                     "graphiques reprises de l'application), mais ce thème est un paramètre du "
                     "serveur de tableaux de bord, fixé à son lancement : il ne peut donc pas "
                     "suivre le bouton clair/sombre du navigateur, qui ne concerne que le chat.")

    add_bullet(doc, "La toute première requête d'un widget prend environ 12 secondes (démarrage à "
                     "froid du moteur de requête), puis ~400 ms. Sensible uniquement au premier "
                     "chargement après démarrage.")
    add_bullet(doc, "Le chargement retraite chaque ligne du Sheet à chaque cycle plutôt que de ne "
                     "toucher que les lignes modifiées — sans impact réel à l'échelle actuelle "
                     "(~360 lignes toutes les 15 minutes), à reconsidérer si le Sheet grossissait "
                     "beaucoup.")
    add_bullet(doc, "Les retouches appliquées par le code ne couvrent pas les valeurs "
                     "dynamiques : « en camembert » ou « par practice » sont traités sans appel au "
                     "modèle, mais « et pour le Maroc ? » repart par le chemin complet, la liste "
                     "des pays venant des données. Délibéré — mieux vaut un appel de plus qu'un "
                     "filtre pays silencieusement ignoré (section 9.4).")
    add_bullet(doc, "Le tableau de bord de travail est un fichier unique côté serveur : deux "
                     "personnes utilisant l'application en même temps réécriraient le même. Sans "
                     "effet en usage local mono-utilisateur, bloquant pour un déploiement "
                     "multi-utilisateurs — qui demanderait de toute façon une authentification.")
    add_bullet(doc, "Le mot « urgentes » seul ne pose aucun filtre de délai : il faut « urgentes "
                     "(< 7 jours) » pour borner l'échéance. Comportement volontaire — ne rien "
                     "deviner — mais qui surprend au premier essai.")

    out_path = os.path.join(DOC_DIR, "Guide_Technique_DevoTeam_Dashboard.docx")
    doc.save(out_path)
    return out_path


# ---------- Document 3 : script de présentation orale ----------

def build_presentation_script():
    doc = new_document()
    add_cover(
        doc,
        "Notes de présentation",
        "DevoTeam Dashboard",
        "Script pour une présentation orale — à lire, pas à distribuer",
    )

    add_h1(doc, "Accroche")
    add_cue(doc, "[Ton confiant, ne pas se presser — laisser la phrase respirer]")
    add_body(doc,
        "Je vais vous présenter DevoTeam Dashboard. En une phrase : c'est un tableau de "
        "bord commercial à qui on parle en français, et qui répond avec le bon graphique — "
        "pas de menu à chercher, pas de SQL à écrire, pas de formation à prévoir.")
    add_body(doc,
        "L'idée de départ tient en une ligne : « montre-moi le budget par pays pour Risk "
        "Advisory », et une seconde plus tard, le graphique est là, à l'écran, à côté du "
        "chat.")

    add_h1(doc, "Le problème que ça résout")
    add_body(doc,
        "Les équipes commerciales suivent des centaines d'opportunités — budgets, pays, "
        "statuts, échéances. Deux frictions reviennent tout le temps : il faut un outil BI "
        "ou des requêtes SQL pour croiser ces données, et une échéance qui approche peut "
        "passer inaperçue si personne ne pense à consulter le dashboard ce jour-là. Ce "
        "projet attaque les deux en même temps : l'accès aux données devient conversationnel, "
        "et le système prévient tout seul quand une échéance approche.")

    add_h1(doc, "Ce qui va vous surprendre : un tableau de bord entier par question")
    add_cue(doc, "[Montrer le chat à l'écran — poser une vraie question en direct si possible]")
    add_body(doc,
        "Quand je pose une question, je ne reçois pas un graphique : je reçois un tableau "
        "de bord entier. Les totaux du périmètre, le graphique qui répond, le même chiffre "
        "sous un autre angle, l'état du pipeline, et le détail ligne par ligne — tous "
        "filtrés exactement pareil. Autrement dit, ça ne répond pas seulement à la "
        "question posée, ça donne de quoi poser la suivante.")
    add_body(doc,
        "Parmi les visualisations, il y a un entonnoir de vente qui montre où les deals se "
        "perdent dans le pipeline, et un nuage de points qui croise le budget avec la "
        "probabilité de gain — est-ce que les gros budgets ont statistiquement plus de "
        "chances d'être gagnés ? Ce sont des choses qu'on ne trouve pas dans un dashboard "
        "interne standard, et le système choisit tout seul quoi afficher selon la question.")
    add_body(doc,
        "Le chat comprend aussi le contexte : je peux demander « et pour Data Management ? » "
        "juste après une première question, et il comprend que je veux le même graphique, "
        "juste filtré différemment. Il comprend aussi les dates relatives — « ce mois-ci », "
        "« le trimestre dernier » — et les comparaisons — « compare la France et le Maroc ».")

    add_h1(doc, "Ce qu'on ne voit pas : pourquoi on peut lui faire confiance")
    add_body(doc,
        "C'est le point sur lequel j'ai le plus travaillé, et c'est invisible à l'écran. "
        "Un assistant qui invente un chiffre plausible mais faux, c'est pire qu'un assistant "
        "qui ne répond pas — surtout sur des données commerciales. Alors la règle absolue "
        "du projet, c'est : jamais deviner. Si le système n'est pas sûr de ce qu'on lui "
        "demande, il pose une question de clarification, plutôt que d'afficher un résultat "
        "séduisant mais inventé.")
    add_body(doc,
        "Concrètement, ça veut dire que le modèle de langage ne touche jamais directement "
        "la base de données. Il produit uniquement une intention structurée, entièrement "
        "validée contre une liste blanche de colonnes et de valeurs réelles avant qu'une "
        "seule requête SQL soit construite. Et cette architecture est vérifiée par cent "
        "tests automatisés, qui tournent sans base de données réelle ni connexion internet.")

    add_h1(doc, "Une fonctionnalité que personne ne m'a demandée")
    add_body(doc,
        "Au-delà du périmètre initial, j'ai ajouté un système d'alertes deadlines : chaque "
        "jour, l'application envoie automatiquement un email récapitulant les opportunités "
        "dont l'échéance approche, et affiche la même liste dans un bandeau du dashboard. "
        "Et j'ai poussé le soin du détail jusqu'à gérer le cas où le serveur serait éteint "
        "pile à l'heure prévue : au redémarrage, le système vérifie tout seul s'il a manqué "
        "l'envoi du jour et le rattrape, sans jamais envoyer le même email deux fois.")

    add_h1(doc, "Des bugs trouvés en testant pour de vrai, pas juste en croisant les doigts")
    add_cue(doc, "[C'est la partie qui montre la rigueur — ne pas la survoler]")
    add_body(doc,
        "Les tests automatisés, c'est nécessaire, mais ça ne suffit pas. J'ai trouvé un vrai "
        "bug en utilisant l'application moi-même : demander la liste des « opportunités "
        "urgentes » remontait des deadlines dépassées depuis des mois, parce que "
        "mathématiquement, moins sept jours est bien inférieur à sept jours — juste que ça "
        "n'a aucun sens métier. Je l'ai corrigé à trois niveaux différents plutôt qu'un "
        "seul, par prudence : le raisonnement, c'est qu'une instruction donnée à un modèle "
        "de langage n'est jamais une garantie à cent pour cent, donc la correction qui "
        "compte vraiment est celle qui ne dépend pas du modèle pour s'appliquer.")
    add_body(doc,
        "Même logique pour deux graphiques : le texte qui accompagnait l'entonnoir de vente "
        "et la carte de chaleur décrivait plus de données que ce que le graphique affichait "
        "réellement à l'écran. Des tests unitaires bien écrits n'auraient jamais attrapé ça "
        "— il a fallu utiliser l'application comme un vrai utilisateur pour le voir.")

    add_h1(doc, "Le soin du détail")
    add_body(doc,
        "L'interface reprend les couleurs et le logo Devoteam, avec un mode sombre "
        "entièrement lisible — et pas juste esthétique : j'ai trouvé et corrigé deux bugs "
        "de contraste où le texte devenait quasiment invisible en mode sombre, en relisant "
        "le CSS ligne par ligne plutôt qu'en devinant. Et pour que ce soit simple à lancer "
        "au quotidien, j'ai construit un raccourci bureau en un clic qui démarre les trois "
        "services et ouvre directement la page.")
    add_body(doc,
        "J'ai aussi fini par supprimer la base de données. Elle ne servait plus à rien : "
        "les équipes saisissent dans un Google Sheet, et on recopiait ce Sheet en base "
        "toutes les quinze minutes pour le relire juste après. Aujourd'hui l'application "
        "lit le Sheet directement — plus rien à installer, plus rien à administrer, et "
        "tout un pan de code compliqué qui disparaît avec.")

    add_h1(doc, "Les chiffres à retenir")
    add_bullet(doc, "1 question = 1 tableau de bord complet, généré automatiquement.")
    add_bullet(doc, "9 types de visualisations choisis automatiquement selon la question.")
    add_bullet(doc, "238 tests automatisés, aucune dépendance à des données réelles.")
    add_bullet(doc, "21 phases de développement livrées, de bout en bout, en autonomie.")
    add_bullet(doc, "0 base de données à administrer — le Google Sheet est la source de vérité.")
    add_bullet(doc, "0 hallucination tolérée : donnée non fiable = question posée en retour, jamais un chiffre inventé.")

    add_h1(doc, "Pour conclure")
    add_cue(doc, "[Marquer une pause avant cette dernière phrase]")
    add_body(doc,
        "Ce projet, c'est un dashboard qu'on peut vraiment donner à une équipe commerciale "
        "sans formation, avec la rigueur d'ingénierie qui garantit que ce qu'il affiche est "
        "juste — et deux documents détaillés, un rapport et un guide technique, pour qui "
        "veut aller plus loin. Je suis disponible pour une démonstration en direct, ou pour "
        "répondre à vos questions dès maintenant.")

    out_path = os.path.join(DOC_DIR, "Script_Presentation_DevoTeam_Dashboard.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    # build_presentation_script() n'est volontairement PAS appelée ici : son .docx a été
    # retiré du dépôt (il ne servait plus), et le régénérer à chaque exécution le ferait
    # silencieusement réapparaître. La fonction reste disponible et à jour — l'appeler
    # explicitement suffit à reproduire le document si le besoin revient.
    p1 = build_rapport_professionnel()
    p2 = build_guide_technique()
    print("Généré :", p1)
    print("Généré :", p2)
